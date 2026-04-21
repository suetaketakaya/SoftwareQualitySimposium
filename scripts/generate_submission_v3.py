#!/usr/bin/env python3
"""SQiP 2026 提出用 docx v3: abstract-final-v2.md を反映し、OOP拡張(v3.0)を含む
   cairosvg不使用 — matplotlib生成済みPNGを直接参照"""
import os
import zipfile
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE = "/Users/suetaketakaya/SoftwareQualitySymposium"
FIG = os.path.join(BASE, "drafts", "figures")
OUT = os.path.join(BASE, "report", "submission_final_2026.docx")

PNGS = {
    "fig1": os.path.join(FIG, "fig1-overview.png"),
    "fig2": os.path.join(FIG, "fig2-requirements-breakdown.png"),
    "fig3": os.path.join(FIG, "fig3-test-comparison.png"),
    "fig4": os.path.join(FIG, "fig4-execution-results.png"),
    "fig5": os.path.join(FIG, "fig5-quality-gap.png"),
}

for k, path in PNGS.items():
    assert os.path.exists(path), f"PNG not found: {path}"
    size = os.path.getsize(path)
    assert size >= 50000, f"PNG too small (<50KB): {path} ({size} bytes)"
    print(f"  OK: {k} = {size:,} bytes ({path})")


# ========== ヘルパー関数 ==========

def heading(doc, text, lv=2):
    x = doc.add_heading(text, level=lv)
    for r in x.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.size = Pt(13 if lv == 2 else 11)
        r.font.name = "游明朝"
    x.paragraph_format.space_before = Pt(10)
    x.paragraph_format.space_after = Pt(4)


def subheading(doc, text):
    x = doc.add_paragraph()
    r = x.add_run(text)
    r.font.size = Pt(10.5)
    r.bold = True
    r.font.name = "游明朝"
    x.paragraph_format.space_before = Pt(8)
    x.paragraph_format.space_after = Pt(2)
    x.paragraph_format.line_spacing = Pt(16)


def para(doc, text, sz=10, after=Pt(4), bold=False, align=None, italic=False):
    x = doc.add_paragraph()
    r = x.add_run(text)
    r.font.size = Pt(sz)
    r.bold = bold
    r.italic = italic
    r.font.name = "游明朝"
    if align:
        x.alignment = align
    x.paragraph_format.space_after = after
    x.paragraph_format.space_before = Pt(0)
    x.paragraph_format.line_spacing = Pt(16)
    return x


def bold_para(doc, text_parts, sz=10, after=Pt(4)):
    x = doc.add_paragraph()
    for txt, is_bold in text_parts:
        r = x.add_run(txt)
        r.font.size = Pt(sz)
        r.bold = is_bold
        r.font.name = "游明朝"
    x.paragraph_format.space_after = after
    x.paragraph_format.space_before = Pt(0)
    x.paragraph_format.line_spacing = Pt(16)
    return x


def code_block(doc, text):
    x = doc.add_paragraph()
    r = x.add_run(text)
    r.font.size = Pt(8.5)
    r.font.name = "Consolas"
    x.paragraph_format.space_after = Pt(4)
    x.paragraph_format.space_before = Pt(2)
    x.paragraph_format.left_indent = Cm(0.8)
    x.paragraph_format.line_spacing = Pt(13)


def add_table(doc, headers, rows, col_widths=None, fs=8.5):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hdr in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = hdr
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(fs)
                r.font.name = "游明朝"
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            c.text = str(val)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(fs)
                    r.font.name = "游明朝"
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()


def figure(doc, png_path, caption, width=Cm(14)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(png_path, width=width)
    p.paragraph_format.space_after = Pt(2)
    c = doc.add_paragraph(caption)
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in c.runs:
        r.font.size = Pt(9)
        r.italic = True
        r.font.name = "游明朝"
    c.paragraph_format.space_after = Pt(6)


# ========== メイン処理 ==========

def main():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    st = doc.styles["Normal"]
    st.font.name = "游明朝"
    st.font.size = Pt(10)
    st.paragraph_format.space_after = Pt(4)
    st.paragraph_format.line_spacing = Pt(16)

    # ======================================================================
    # PAGE 1: 発表申込書
    # ======================================================================
    para(doc, "ソフトウェア品質シンポジウム2026　「経験論文」「経験発表」",
         bold=True, sz=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "発表申込／アブストラクト　記入用紙",
         bold=True, sz=11, align=WD_ALIGN_PARAGRAPH.CENTER, after=Pt(14))

    para(doc, "〔発表申込書〕", bold=True, sz=11, after=Pt(8))

    para(doc, "タイトル", bold=True, sz=10)
    para(doc, "静的解析に基づくテスト要求モデルの自動生成と、"
         "それを用いたLLMテスト生成の実証評価", sz=10.5)

    para(doc, "申込区分", bold=True, sz=10)
    para(doc, "経験発表", sz=10.5)

    para(doc, "カテゴリ", bold=True, sz=10)
    para(doc, "品質管理・テスト技術の観点", sz=10.5)

    para(doc, "キーワード", bold=True, sz=10)
    para(doc, "ホワイトボックステスト / テスト要求モデル / LLM / "
         "静的解析 / 単体テスト生成 / トレーサビリティ", sz=10.5, after=Pt(14))

    # ======================================================================
    # PAGE 2+: アブストラクト
    # ======================================================================
    doc.add_page_break()
    para(doc, "〔アブストラクト記入用紙〕", bold=True, sz=11,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=Pt(10))

    para(doc,
         "静的解析に基づくテスト要求モデルの自動生成と、"
         "それを用いたLLMテスト生成の実証評価",
         bold=True, sz=13, align=WD_ALIGN_PARAGRAPH.CENTER, after=Pt(4))

    para(doc, "申込区分: 経験発表", sz=10, after=Pt(2))
    para(doc, "カテゴリ: 品質管理・テスト技術の観点", sz=10, after=Pt(2))
    para(doc, "キーワード: ホワイトボックステスト / テスト要求モデル / LLM / "
         "静的解析 / 単体テスト生成 / トレーサビリティ", sz=10, after=Pt(8))

    # ------------------------------------------------------------------
    # 1. ねらい
    # ------------------------------------------------------------------
    heading(doc, "1. ねらい")

    subheading(doc, "1.1 背景")

    para(doc,
         "大規模言語モデル（LLM）の急速な発展により、ソフトウェア開発のワークフローは"
         "大きく変化しつつある。現在のLLM活用開発の主流は、荒い要求仕様からLLMに実装コードを"
         "直接生成させるパターンであり、評価が「要求から現れる基本機能の動作確認」に偏る傾向がある。"
         "しかし、LLMが生成したコードを実リリースに載せるためには、"
         "単体テスト・コンポーネントテストの層で品質を担保する必要がある。")

    para(doc,
         "この傾向は研究コミュニティにおいても再現されている。"
         "標準ベンチマークであるHumanEval [18] およびMBPP [19] は"
         "pass@kのみを評価指標としており、テストカバレッジや非機能品質は評価対象外である。"
         "Liu et al. [20] は、テストケースを80倍に拡張するとpass@kが19.3〜28.9%低下することを報告した。")

    para(doc,
         "現状のLLMテスト生成手法（CodaMOSA [1]、ChatUniTest [2]、TestPilot [3]、"
         "CoverUp [4]、CoverAgent [5]、TELPA [6] 等）はカバレッジ向上やコンパイル成功率の改善に"
         "注力している。Chu et al. [7] による115本の体系的サーベイでは、LLMテスト生成研究の89%が"
         "「プロンプトエンジニアリングによる直接テスト生成」方式を採り、テスト設計の意図は"
         "LLM内部に閉じたまま外部化されず、以下の3つの問題が構造的に未解決である。")

    subheading(doc, "1.2 問題定義")

    bold_para(doc, [
        ("(P1) テスト設計根拠の不在: ", True),
        ("LLMが生成したテストが「何の分岐を」「どの同値クラスを」対象としているか不明であり、"
         "テスト設計の妥当性を第三者が検証できない。", False),
    ])

    bold_para(doc, [
        ("(P2) 網羅性の事前把握が不可能: ", True),
        ("テスト群がソースコードのどの範囲をカバーしているかは実行するまで分からない。"
         "テスト生成「前」に網羅性を設計・管理する手段がない。", False),
    ])

    bold_para(doc, [
        ("(P3) テスト失敗時の原因特定コスト: ", True),
        ("テスト失敗時に、生成テストのアサーションとソースコードを逐一突合する必要があり、"
         "原因分類に時間がかかる。", False),
    ])

    subheading(doc, "1.3 目的と提案")

    para(doc,
         "本研究は、これら3つの問題に対し、ソースコードの静的解析結果から"
         "ホワイトボックステスト要求モデル（以下TRM: Test Requirement Model）をYAML形式で自動生成し、"
         "TRMを入力としてLLMにテストコードを生成させる手法を提案する。"
         "TRMは分岐条件・同値クラス・境界値をID付きで構造化した中間成果物であり、"
         "テスト生成の「前」にテスト設計の網羅性を定量評価できる。"
         "さらに、TRMに対する独立した網羅性監査と、監査結果に基づくテスト設計書の生成を行い、"
         "C0/C1/MC-DCカバレッジの達成を検証する。"
         "OSSプロジェクトのC++コード8関数に適用し、TRMの生成精度、網羅性監査プロセスの有効性、"
         "テスト生成品質、テスト失敗時の診断効率を定量評価した。")

    subheading(doc, "1.4 先行研究との差異")

    para(doc,
         "先行手法はいずれも「LLMに直接テストコードを生成させ、その精度を反復的に改善する」"
         "アーキテクチャを採る [7]。Mathews & Nagappan [8] は、CoverAgent等が"
         "「テスト通過」を優先する設計によりバグ発見と矛盾することを指摘している。"
         "また、LLM生成テストにはテストスメルが頻出し [9]、"
         "テストアサーションの正確性は50%以下であることが報告されている [10]。")

    para(doc,
         "本研究はテスト生成の精度向上ではなく、テスト設計情報の構造化・外部化に焦点を当てる点で"
         "既存手法と目的が異なる。BDD/Specification by Example [11] は思想を共有するが"
         "受入テスト層が対象であり、ACL 2025の多段階テスト仕様生成 [12] はトップダウン方式で"
         "本研究のボトムアップ方式とは方向が異なる。")

    subheading(doc, "1.5 本研究の新規性")

    bold_para(doc, [
        ("新規性1: TRMという構造化中間成果物の導入。", True),
        ("YAML形式のTRMを導入し、テスト設計の意図をLLM内部から外部化した。"
         "TRMは5種別（BR/EC/BV/ER/DP）でテスト要求を分類し、各要求にIDを付与して"
         "テストコードへのトレーサビリティを維持する。", False),
    ])

    bold_para(doc, [
        ("新規性2: LLM生成テスト設計に対する網羅性監査プロセス。", True),
        ("初回TRM（99件）の漏れ64件を特定して163件に拡充した"
         "（分岐網羅率87%→推定97%）。", False),
    ])

    bold_para(doc, [
        ("新規性3: テスト品質の種別別可視化。", True),
        ("既存テストの品質ギャップを5種別で定量化し、"
         "テスト品質の構造的な偏りを可視化した。", False),
    ])

    # ------------------------------------------------------------------
    # 2. 実施概要
    # ------------------------------------------------------------------
    heading(doc, "2. 実施概要")

    subheading(doc, "2.1 手法")

    para(doc,
         "提案手法は5段階のパイプラインである（図1）。各段階において大規模言語モデルを活用する。")

    bold_para(doc, [
        ("Step 1. 対象選定（静的解析 + LLM推論）: ", True),
        ("対象リポジトリのソースコード構造を解析し、テスト対象関数を選定する。選定基準は "
         "(a) 外部依存なし（純粋関数）、(b) 分岐構造が明確、(c) 既存テストが存在（比較用）。", False),
    ])

    bold_para(doc, [
        ("Step 2. TRM生成（静的解析 + LLM推論）: ", True),
        ("選定した各関数のソースコードから分岐条件を抽出し、5種別でテスト要求を定義する（表1）。"
         "BR（分岐網羅）はAmmann & Offutt [24] のグラフカバレッジ基準に、"
         "EC/BVはMyers [25] の同値分割法・境界値分析法に、"
         "ERはISO 29119-4 [26] のエラー推測にそれぞれ対応する。", False),
    ])

    para(doc, "表1: TRM種別定義（関数レベル5種別）", bold=True, sz=9,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=Pt(2))
    add_table(doc,
              ["種別", "記号", "定義", "理論的根拠"],
              [
                  ["分岐網羅", "BR", "if/else/switchの各パスを検証", "Ammann & Offutt [24]"],
                  ["同値クラス", "EC", "入力パラメータの同値分割を検証", "Myers [25]"],
                  ["境界値", "BV", "同値クラス境界の上下限を検証", "Myers [25]"],
                  ["エラーパス", "ER", "異常系・エラー処理を検証", "ISO 29119-4 [26]"],
                  ["依存切替", "DP", "関数間の依存関係・差異を検証", "関数間結合テスト"],
              ],
              col_widths=[2.0, 1.0, 5.0, 3.5], fs=8.5)

    para(doc,
         "各要求には一意のID、自然言語の説明、優先度、根拠となるソースコード行番号を付与し、"
         "YAML形式で出力する。以下はParseVersion関数に対するTRMの抜粋である。")

    code_block(doc,
               '- id: "BR-11"\n'
               '  type: branch_coverage\n'
               '  description: "alpha修飾子の分岐を通過する"\n'
               '  priority: high\n'
               '  source_ref: "BC-02-10: *p == \'a\' && strncmp == \'alpha\'"\n'
               '  expected_behavior: "nShift = -0x60 を設定"')

    bold_para(doc, [
        ("Step 3. TRM網羅性監査（静的解析 + 独立レビュー）: ", True),
        ("生成されたTRMに対し、ソースコードの全分岐条件をTRM要求と1対1で照合し、"
         "漏れている条件を特定する。監査結果に基づき追加要求をTRMに統合し、"
         "関数ごとのテスト設計書を生成する。", False),
    ])

    bold_para(doc, [
        ("Step 4. テストコード生成（LLM）: ", True),
        ("監査済みTRMのYAML全体と対象ソースコードをLLMに入力し、テストコードを出力させる。"
         "各テストケースにはTRM IDをコメントで記載し、要求からテストへの追跡性を確保する。", False),
    ])

    code_block(doc,
               'TEST(ParseVersion, AlphaModifier) {\n'
               '    // BR-11: alpha修飾子の分岐を通過する\n'
               '    UINT32 val = ParseVersion(L"2.4.1alpha");\n'
               '    EXPECT_EQ(val, static_cast<UINT32>(0x82848120));\n'
               '}')

    bold_para(doc, [
        ("Step 5. カバレッジ精密分析（静的解析）: ", True),
        ("全テストスイートに対し、C0/C1/MC-DC/ループ境界でカバレッジを精密分析する。", False),
    ])

    figure(doc, PNGS["fig1"],
           "図1: 提案手法の全体構成（5段階パイプライン）", Cm(14.5))

    subheading(doc, "2.2 実験条件")

    para(doc, "表2: 実験条件と対象関数", bold=True, sz=9,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=Pt(2))
    add_table(doc,
              ["領域", "対象関数", "行数", "既存テスト"],
              [
                  ["日時書式・バージョン解析\n(format.cpp)",
                   "GetDateTimeFormat,\nParseVersion,\nCompareVersion",
                   "約115行", "約20件"],
                  ["メールアドレス・文字種判定\n(CWordParse.cpp)",
                   "IsMailAddress,\nWhatKindOfTwoChars,\nWhatKindOfTwoChars4KW",
                   "約180行", "約15件"],
                  ["全角半角英数変換\n(convert_util.cpp)",
                   "Convert_ZeneisuToHaneisu,\nConvert_HaneisuToZeneisu",
                   "約90行", "約10件"],
                  ["合計", "8関数", "約385行", "約45件"],
              ],
              col_widths=[3.5, 4.0, 1.5, 1.5], fs=8.5)

    para(doc,
         "対象: sakura-editor/sakura（C++、OSSテキストエディタ、Google Test整備済み）。"
         "環境: macOS (Apple clang 16.0) + Google Test 1.17。"
         "使用モデル: 大規模言語モデル（匿名査読のためモデル名は伏せる）。"
         "比較対象: (A) 既存の手動テスト（約45件）、(C) 提案手法によるTRM + 生成テスト。")

    # ------------------------------------------------------------------
    # 3. 実施結果
    # ------------------------------------------------------------------
    heading(doc, "3. 実施結果")

    subheading(doc, "3.1 TRM生成結果（P1に対応）")

    para(doc,
         "8関数に対し計99件のテスト要求を初回生成した。種別ごとの内訳を表3に示す（図2）。")

    para(doc, "表3: TRM種別別内訳（初回生成）", bold=True, sz=9,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=Pt(2))
    add_table(doc,
              ["種別", "件数", "割合", "行数/要求"],
              [
                  ["BR（分岐網羅）", "55", "55.6%", "7.0"],
                  ["EC（同値クラス）", "27", "27.3%", "14.3"],
                  ["BV（境界値）", "11", "11.1%", "35.0"],
                  ["ER（エラーパス）", "3", "3.0%", "128.3"],
                  ["DP（依存切替）", "3", "3.0%", "128.3"],
                  ["合計", "99", "100%", "3.9"],
              ],
              col_widths=[3.5, 1.5, 1.5, 2.0], fs=9)

    para(doc,
         "BR要求が最多であるのは、ソースコードの分岐から直接導出されるためである。"
         "ER/DPが少ないのは対象が純粋関数であり、エラーハンドリングや関数間依存が限定的であることによる。"
         "分岐密度の高い関数ほどテスト要求密度が高く、TRMがコードの複雑度に比例した適切な粒度で"
         "品質要求を抽出していることを示す。"
         "可読性分類では、L1+L2（コード知識不要）が65件（65.7%）であった。")

    figure(doc, PNGS["fig2"],
           "図2: テスト要求の種別内訳（N=99）", Cm(10))

    # 3.2 TRM網羅性監査
    subheading(doc, "3.2 TRM網羅性監査結果（新規性2に対応）")

    para(doc,
         "初回TRM（99件）に対し独立監査を実施した。ソースコードの全分岐条件（91件）と"
         "TRMのBR要求を1対1で照合した結果、初回の分岐網羅率は87%（79/91）であった。")

    para(doc, "表4: 監査による追加要求の内訳", bold=True, sz=9,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=Pt(2))
    add_table(doc,
              ["カテゴリ", "初回", "追加", "監査後合計"],
              [
                  ["BR（分岐網羅）", "55", "10", "65"],
                  ["EC（同値クラス）", "27", "21", "48"],
                  ["BV（境界値）", "11", "18", "29"],
                  ["ER（エラーパス）", "3", "10", "13"],
                  ["DP（依存切替）", "3", "5", "8"],
                  ["合計", "99", "64", "163"],
              ],
              col_widths=[3.5, 2.0, 2.0, 2.5], fs=9)

    para(doc,
         "追加要求の重要度はCritical 4件、High 11件、Medium 36件、Low 13件であった。"
         "Critical指摘にはバッファオーバーフロー、intキャスト時の符号反転、"
         "size_tラップアラウンドが含まれる。"
         "監査後の推定分岐網羅率は97%に改善された。"
         "境界値（11→29件、2.6倍増）やエラーパス（3→13件、4.3倍増）の抽出に"
         "LLMの構造的弱点があることを定量的に示した。")

    # 3.3 テスト生成
    subheading(doc, "3.3 テスト生成結果")

    para(doc,
         "監査済みTRM（163件）を入力として計248件のテストコードを生成した。"
         "初回TRM対応145件＋追加要求対応103件で構成される。"
         "カバー率100%。既存テスト約45件に対し約5.5倍を体系的に生成した（図3）。")

    figure(doc, PNGS["fig3"],
           "図3: 既存テストと生成テストの領域別比較", Cm(12))

    # 3.4 実行結果
    subheading(doc, "3.4 コンパイル・実行結果")

    add_table(doc,
              ["指標", "値"],
              [
                  ["初回テストPASS", "133/145件 (91.7%)"],
                  ["初回テストFAIL", "12件 (8.3%)"],
                  ["FAIL原因: 対象関数のバグ", "0件"],
                  ["FAIL原因: LLMの期待値推論誤り", "12件 (100%)"],
                  ["修正後テストPASS", "145/145件 (100%)"],
                  ["全248件最終PASS", "248/248件 (100%)"],
              ],
              col_widths=[5.0, 5.0], fs=9)

    figure(doc, PNGS["fig4"],
           "図4: コンパイル・実行結果（初回→修正後）", Cm(12))

    # 3.5 FAIL分析
    subheading(doc, "3.5 FAIL分析とTRMの診断効果（P3に対応）")

    add_table(doc,
              ["原因分類", "件数", "対応TRM ID", "技術的詳細"],
              [
                  ["コンポーネント分割誤解", "5", "BR-11〜15",
                   '"2.4.1alpha"を1コンポーネントと誤予測'],
                  ["数字グルーピング誤解", "1", "EC内部",
                   '"1234"を1桁ずつと誤予測'],
                  ["文字数カウントミス", "2", "BR-34, EC",
                   "wcslen=18をnull込み19と誤算"],
                  ["API呼び出しバグ", "1", "EC-16",
                   "offsetとポインタの二重シフト"],
                  ["マッピング戻り値誤解", "3", "BR-39〜41",
                   "変数マッピング後の返却値を誤認"],
              ],
              col_widths=[3.0, 1.0, 2.5, 5.5], fs=8.5)

    para(doc,
         "TRM IDから該当要求の日本語説明と根拠を即座に特定でき、原因分類が構造的に行えた。"
         "12件全てが要求レベルで分類可能であり、うち約8件はコード非専門者にも診断可能であった。")

    # 3.6 品質ギャップ
    subheading(doc, "3.6 品質ギャップ分析（P2に対応）")

    add_table(doc,
              ["種別", "TRM要求数", "カバー済み（推定）", "未カバー", "未カバー率"],
              [
                  ["BR", "55", "26", "29", "52.7%"],
                  ["EC", "27", "8", "19", "70.4%"],
                  ["BV", "11", "3", "8", "72.7%"],
                  ["ER", "3", "1", "2", "66.7%"],
                  ["DP", "3", "0", "3", "100%"],
                  ["合計", "99", "38", "61", "61.6%"],
              ],
              col_widths=[1.8, 2.0, 2.5, 2.0, 2.0], fs=8.5)

    para(doc,
         "DP(100%) > BV(72.7%) > EC(70.4%) > ER(66.7%) > BR(52.7%)"
         "の順で未カバー率が高い。既存テストが分岐の主要パス（BR）に偏っていることが示された。"
         "この分析はTRMが存在するからこそ可能であり、P2への解決策として有効に機能する。")

    figure(doc, PNGS["fig5"],
           "図5: 既存テストの品質ギャップ分析（種別別未カバー率）", Cm(10))

    # 3.7 カバレッジ
    subheading(doc, "3.7 カバレッジ精密分析結果")

    para(doc, "表6: カバレッジ達成状況", bold=True, sz=9,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=Pt(2))
    add_table(doc,
              ["基準", "達成値", "備考"],
              [
                  ["C0（ステートメント）", "155/155 = 100%", "全8関数で100%"],
                  ["C1（ブランチ）", "146/146 = 100%", "全8関数で100%"],
                  ["MC/DC", "83/84 = 99%", "未カバー1件はIsMailAddressの補完的条件"],
                  ["ループ境界", "33/33 = 100%", "0回/1回/N回の全パターン"],
              ],
              col_widths=[3.0, 3.0, 5.5], fs=9)

    para(doc,
         "MC/DCは航空宇宙（DO-178C [27]）や自動車（ISO 26262 [28]）で要求される高水準の基準であり、"
         "LLMテスト生成とTRMの組み合わせがこの水準を達成可能であることを示した。")

    # ------------------------------------------------------------------
    # 4. 結論
    # ------------------------------------------------------------------
    heading(doc, "4. 結論")

    subheading(doc, "4.1 まとめ")

    para(doc,
         "本研究では、ソースコードの静的解析からTRMをYAML形式で生成し、"
         "TRM網羅性監査を経てテスト設計書を作成し、TRMを入力としてLLMに"
         "単体テストを生成させる手法を提案・実証した。C++コード8関数（約385行）に適用した結果:")

    bold_para(doc, [
        ("1. TRM生成と網羅性監査: ", True),
        ("5種別99件を初回生成し、監査により64件追加して163件に拡充。"
         "分岐網羅率87%→97%。可読率65.7%。", False),
    ])

    bold_para(doc, [
        ("2. テスト設計書の生成: ", True),
        ("関数ごとの正式なテスト設計書を生成。"
         "テスト設計レビューとトレーサビリティの基盤として機能。", False),
    ])

    bold_para(doc, [
        ("3. カバレッジ達成: ", True),
        ("248件のテストを生成。C0/C1 100%、MC/DC 99%、ループ境界100%。", False),
    ])

    bold_para(doc, [
        ("4. 実行精度: ", True),
        ("初回PASS率91.7%。FAIL 12件は全てLLM期待値誤り。修正後100% PASS。", False),
    ])

    bold_para(doc, [
        ("5. 品質ギャップ発見: ", True),
        ("既存テストの未カバー要求61件（61.6%）を種別ごとに特定。", False),
    ])

    para(doc,
         "TRMの導入により、テスト設計根拠の明示（P1）、網羅性の事前定量管理（P2）、"
         "テスト失敗時の構造的原因分類（P3）を実現した。")

    subheading(doc, "4.2 今後の課題")

    bold_para(doc, [
        ("(1) OOP構造への拡張: ", True),
        ("本実験は純粋関数に限定されており、クラス継承やインスタンス変数の状態管理は対象外であった。"
         "実務ではクラス継承階層、仮想メソッドの多態性、インスタンス変数の状態遷移、"
         "言語固有イディオム（RAII、コンテキストマネージャ等）がテスト品質に大きく影響する。"
         "TRMスキーマにCI（クラス継承: 多態性・オーバーライド・リスコフ置換）、"
         "SV（状態変数: 初期化・状態遷移・不変条件維持）、"
         "CP（コードパターン: デザインパターン・言語イディオム・フレームワーク規約）の"
         "3種別を追加し、8種別体系への拡張を進めている。", False),
    ])

    bold_para(doc, [
        ("(2) TRM網羅性監査の自動化: ", True),
        ("gcov/llvm-cov等のカバレッジツール出力とTRM要求の自動照合により、"
         "監査プロセスの効率化と客観性の向上が期待できる。", False),
    ])

    bold_para(doc, [
        ("(3) 可読性の被験者評価: ", True),
        ("可読率65.7%は著者による分類であり、非開発者による理解度評価実験が必要である。", False),
    ])

    bold_para(doc, [
        ("(4) 他言語・他プロジェクトでの再現: ", True),
        ("TRM種別体系は言語非依存に設計されており、Python/Java/TypeScript/Go/Rust等での"
         "再現実験により一般化可能性を検証する。パイプラインは任意のGitHubリポジトリを入力として"
         "受け付ける汎用ツールとして実装済みである。", False),
    ])

    # ------------------------------------------------------------------
    # 参考文献
    # ------------------------------------------------------------------
    para(doc, "", sz=4, after=Pt(2))
    heading(doc, "参考文献")

    refs = [
        '[1] C. Lemieux et al., "CodaMOSA," Proc. ICSE 2023, pp. 919-931.',
        '[2] Y. Chen et al., "ChatUniTest," Companion Proc. FSE 2024, ACM.',
        '[3] M. Schafer et al., "LLMs for Automated Unit Test Generation," IEEE TSE, vol. 50, 2024.',
        '[4] J. A. Ramos et al., "CoverUp," arXiv:2403.16218, 2024.',
        '[5] Qodo, "Qodo-Cover," GitHub, 2024. https://github.com/Codium-ai/cover-agent',
        '[6] "TELPA," ACM TOSEM, 2024. https://dl.acm.org/doi/10.1145/3748505',
        '[7] B. Chu et al., "LLMs for Unit Test Generation: Survey," arXiv:2511.21382, 2025.',
        '[8] N. S. Mathews and M. Nagappan, arXiv:2412.14137, 2024.',
        '[9] "Test smells in LLM-Generated Unit Tests," arXiv:2410.10628, 2024.',
        '[10] "Do LLMs generate test oracles...?" arXiv:2410.21136, 2024.',
        '[11] G. Adzic, Specification by Example, Manning, 2011.',
        '[12] "Multi-Step Generation of Test Specifications using LLMs," ACL 2025 Industry.',
        '[13] C. Wiecher et al., "Model-based analysis," Systems Engineering, Wiley, 2024.',
        '[14] "Static Analysis as a Feedback Loop," arXiv:2508.14419, 2025.',
        '[15] 「テストパターンマトリックスを用いたテスト生成」, SQiP 2024 A4-1.',
        '[16] AIST, 「生成AI品質マネジメントガイドライン第1版」, 2025.',
        '[17] IPA, 「テキスト生成AIの導入・運用ガイドライン」, 2024.',
        '[18] M. Chen et al., "Evaluating LLMs Trained on Code," arXiv:2107.03374, 2021.',
        '[19] J. Austin et al., "Program Synthesis with LLMs," arXiv:2108.07732, 2021.',
        '[20] J. Liu et al., "Is Your Code Generated by ChatGPT Really Correct?" NeurIPS 2023.',
        '[21] J. Zheng et al., "Beyond Correctness," arXiv:2407.11470, 2024.',
        '[22] N. Perry et al., "Do Users Write More Insecure Code with AI?" ACM CCS 2023.',
        '[23] SonarSource, "AI Code Assurance," 2026.',
        '[24] P. Ammann and J. Offutt, Introduction to Software Testing, 2nd ed., CUP, 2016.',
        '[25] G.J. Myers et al., The Art of Software Testing, 3rd ed., Wiley, 2012.',
        '[26] ISO/IEC/IEEE 29119-4:2021, "Software Testing -- Part 4: Test Techniques."',
        '[27] RTCA, DO-178C, 2011.',
        '[28] ISO 26262:2018, "Road Vehicles -- Functional Safety, Part 6."',
        '[29] C. Cadar and K. Sen, "Symbolic Execution," CACM, vol. 56, no. 2, 2013.',
        '[30] M. Utting and B. Legeard, Practical Model-Based Testing, Morgan Kaufmann, 2007.',
    ]

    for ref in refs:
        para(doc, ref, sz=8, after=Pt(2))

    # ========== 保存 ==========
    doc.save(OUT)
    print(f"\nSaved: {OUT}")
    print(f"File size: {os.path.getsize(OUT):,} bytes")

    print("\n--- 埋め込み画像サイズの検証 ---")
    with zipfile.ZipFile(OUT, 'r') as z:
        image_files = [n for n in z.namelist() if n.startswith('word/media/')]
        assert len(image_files) == 5, f"Expected 5 embedded images, found {len(image_files)}"
        for img_name in sorted(image_files):
            info = z.getinfo(img_name)
            size_bytes = info.file_size
            print(f"  {img_name}: {size_bytes:,} bytes")
            assert size_bytes >= 70000, (
                f"Embedded image too small (<70KB): {img_name} ({size_bytes:,} bytes)"
            )
    print("\nAll 5 embedded images are >= 70KB. Verification passed.")


if __name__ == "__main__":
    main()
