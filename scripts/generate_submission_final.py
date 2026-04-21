#!/usr/bin/env python3
"""SQiP 2026 最終提出用 docx: abstract-final.md を忠実にdocx化（cairosvg不使用）"""
import os
import zipfile
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE = "/Users/suetaketakaya/SoftwareQualitySymposium"
FIG = os.path.join(BASE, "drafts", "figures")
OUT = os.path.join(BASE, "report", "submission_final_2026.docx")

# matplotlib生成済みPNGのパスを直接参照（cairosvg不使用・SVG変換一切なし）
PNGS = {
    "fig1": os.path.join(FIG, "fig1-overview.png"),
    "fig2": os.path.join(FIG, "fig2-requirements-breakdown.png"),
    "fig3": os.path.join(FIG, "fig3-test-comparison.png"),
    "fig4": os.path.join(FIG, "fig4-execution-results.png"),
    "fig5": os.path.join(FIG, "fig5-quality-gap.png"),
}

# ========== PNGファイルの存在確認・サイズ検証 ==========
for k, path in PNGS.items():
    assert os.path.exists(path), f"PNG not found: {path}"
    size = os.path.getsize(path)
    assert size >= 50000, f"PNG too small (<50KB): {path} ({size} bytes)"
    print(f"  OK: {k} = {size:,} bytes ({path})")


# ========== ヘルパー関数 ==========

def heading(doc, text, lv=2):
    """セクション見出し（## レベル）"""
    x = doc.add_heading(text, level=lv)
    for r in x.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.size = Pt(13 if lv == 2 else 11)
        r.font.name = "游明朝"
    x.paragraph_format.space_before = Pt(10)
    x.paragraph_format.space_after = Pt(4)


def subheading(doc, text):
    """サブセクション見出し（### レベル）"""
    x = doc.add_paragraph()
    r = x.add_run(text)
    r.font.size = Pt(10.5)
    r.bold = True
    r.font.name = "游明朝"
    x.paragraph_format.space_before = Pt(8)
    x.paragraph_format.space_after = Pt(2)
    x.paragraph_format.line_spacing = Pt(16)


def para(doc, text, sz=10, after=Pt(4), bold=False, align=None, italic=False):
    """本文段落"""
    x = doc.add_paragraph()
    r = x.add_run(text)
    r.font.size = Pt(sz)
    r.bold = bold
    r.italic = italic
    r.font.name = "游明朝"
    if align:
        x.alignment = align
    x.paragraph_format.space_after = after
    x.paragraph_format.space_before = Pt(0)
    x.paragraph_format.line_spacing = Pt(16)
    return x


def bold_para(doc, text_parts, sz=10, after=Pt(4)):
    """太字と通常テキストの混在段落。text_partsは(text, bold)のリスト"""
    x = doc.add_paragraph()
    for txt, is_bold in text_parts:
        r = x.add_run(txt)
        r.font.size = Pt(sz)
        r.bold = is_bold
        r.font.name = "游明朝"
    x.paragraph_format.space_after = after
    x.paragraph_format.space_before = Pt(0)
    x.paragraph_format.line_spacing = Pt(16)
    return x


def code_block(doc, text):
    """コードブロック（Consolas 8.5pt）"""
    x = doc.add_paragraph()
    r = x.add_run(text)
    r.font.size = Pt(8.5)
    r.font.name = "Consolas"
    x.paragraph_format.space_after = Pt(4)
    x.paragraph_format.space_before = Pt(2)
    x.paragraph_format.left_indent = Cm(0.8)
    x.paragraph_format.line_spacing = Pt(13)


def add_table(doc, headers, rows, col_widths=None, fs=8.5):
    """表の追加"""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hdr in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = hdr
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(fs)
                r.font.name = "游明朝"
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            c.text = str(val)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(fs)
                    r.font.name = "游明朝"
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()  # 表後の余白


def figure(doc, png_path, caption, width=Cm(14)):
    """図の挿入（PNG直接参照）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(png_path, width=width)
    p.paragraph_format.space_after = Pt(2)
    c = doc.add_paragraph(caption)
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in c.runs:
        r.font.size = Pt(9)
        r.italic = True
        r.font.name = "游明朝"
    c.paragraph_format.space_after = Pt(6)


# ========== メイン処理 ==========

def main():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    # 基本スタイル設定
    st = doc.styles["Normal"]
    st.font.name = "游明朝"
    st.font.size = Pt(10)
    st.paragraph_format.space_after = Pt(4)
    st.paragraph_format.line_spacing = Pt(16)

    # ======================================================================
    # PAGE 1: 発表申込書
    # ======================================================================
    para(doc, "ソフトウェア品質シンポジウム2026　「経験論文」「経験発表」",
         bold=True, sz=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "発表申込／アブストラクト　記入用紙",
         bold=True, sz=11, align=WD_ALIGN_PARAGRAPH.CENTER, after=Pt(14))

    para(doc, "〔発表申込書〕", bold=True, sz=11, after=Pt(8))

    para(doc, "タイトル", bold=True, sz=10)
    para(doc, "静的解析に基づくテスト要求モデルの自動生成と、"
         "それを用いたLLMテスト生成の実証評価", sz=10.5)

    para(doc, "申込区分", bold=True, sz=10)
    para(doc, "経験発表", sz=10.5)

    para(doc, "カテゴリ", bold=True, sz=10)
    para(doc, "品質管理・テスト技術の観点", sz=10.5)

    para(doc, "キーワード", bold=True, sz=10)
    para(doc, "ホワイトボックステスト / テスト要求モデル / LLM / "
         "静的解析 / 単体テスト生成 / トレーサビリティ", sz=10.5, after=Pt(14))

    # ======================================================================
    # PAGE 2+: アブストラクト（abstract-final.mdの忠実な変換）
    # ======================================================================
    doc.add_page_break()
    para(doc, "〔アブストラクト記入用紙〕", bold=True, sz=11,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=Pt(10))

    # タイトル
    para(doc,
         "静的解析に基づくテスト要求モデルの自動生成と、"
         "それを用いたLLMテスト生成の実証評価",
         bold=True, sz=13, align=WD_ALIGN_PARAGRAPH.CENTER, after=Pt(4))

    # メタ情報
    para(doc, "申込区分: 経験発表", sz=10, after=Pt(2))
    para(doc, "カテゴリ: 品質管理・テスト技術の観点", sz=10, after=Pt(2))
    para(doc, "キーワード: ホワイトボックステスト / テスト要求モデル / LLM / "
         "静的解析 / 単体テスト生成 / トレーサビリティ", sz=10, after=Pt(8))

    # ------------------------------------------------------------------
    # 1. ねらい
    # ------------------------------------------------------------------
    heading(doc, "1. ねらい")

    # 1.1 背景
    subheading(doc, "1.1 背景")

    para(doc,
         "大規模言語モデル（LLM）の急速な発展により、ソフトウェア開発のワークフローは"
         "大きく変化しつつある。現在のLLM活用開発の主流は、荒い要求仕様からLLMに実装コードを"
         "直接生成させるパターンであり、評価が「要求から現れる基本機能の動作確認」に偏る傾向がある。"
         "しかし、LLMが生成したコードを実リリースに載せるためには、"
         "単体テスト・コンポーネントテストの層で品質を担保する必要がある。")

    para(doc,
         "現状のLLMテスト生成手法（CodaMOSA [1]、ChatUniTest [2]、TestPilot [3]、"
         "CoverUp [4]、CoverAgent [5]、TELPA [6] 等）はカバレッジ向上やコンパイル成功率の改善に"
         "注力している。Chu et al. [7] による115本の体系的サーベイでは、LLMテスト生成研究の89%が"
         "「プロンプトエンジニアリングによる直接テスト生成」方式を採り、反復的な検証・修復ループで"
         "精度を改善するアーキテクチャが標準となっている。このアーキテクチャでは、テスト設計の意図は"
         "LLM内部に閉じたまま外部化されず、以下の3つの問題が構造的に未解決である。")

    # 1.2 問題定義
    subheading(doc, "1.2 問題定義")

    bold_para(doc, [
        ("(P1) テスト設計根拠の不在: ", True),
        ("LLMが生成したテストが「何の分岐を」「どの同値クラスを」対象としているか不明であり、"
         "テスト設計の妥当性を第三者が検証できない。レビュー担当者は生成されたテストの妥当性を"
         "判断できず、結果として「LLMが作ったテストだから信頼できない」という属人的判断に陥る。", False),
    ])

    bold_para(doc, [
        ("(P2) 網羅性の事前把握が不可能: ", True),
        ("テスト群がソースコードのどの範囲をカバーしているかは、コンパイル・実行してカバレッジツールを"
         "通すまで分からない。テスト生成「前」に網羅性を設計・管理する手段がなく、テスト件数が多くても"
         "同じパスの重複テストばかりという状況を事前に検知できない。", False),
    ])

    bold_para(doc, [
        ("(P3) テスト失敗時の原因特定コスト: ", True),
        ("テスト失敗時に、生成テストのアサーションとソースコードを逐一突合する必要があり、"
         "原因分類に時間がかかる。例えば145件のテストが一斉に失敗した場合、"
         "テストコードとソースコードの突合に長時間を要する。", False),
    ])

    # 1.3 目的と提案
    subheading(doc, "1.3 目的と提案")

    para(doc,
         "本研究は、これら3つの問題に対し、ソースコードの静的解析結果から"
         "ホワイトボックステスト要求モデル（以下TRM: Test Requirement Model）をYAML形式で自動生成し、"
         "TRMを入力としてLLMにテストコードを生成させる手法を提案する。"
         "TRMは分岐条件・同値クラス・境界値をID付きで構造化した中間成果物であり、"
         "テスト生成の「前」にテスト設計の網羅性を定量評価できる。"
         "OSSプロジェクトのC++コード8関数に適用し、TRMの生成精度、テスト生成品質、"
         "テスト失敗時の診断効率を定量評価した。")

    # 1.4 先行研究との差異
    subheading(doc, "1.4 先行研究との差異")

    para(doc,
         "先行手法はいずれも「LLMに直接テストコードを生成させ、その精度を反復的に改善する」"
         "アーキテクチャを採る [7]。Mathews & Nagappan [8] は、CoverAgent等のLLMテスト生成ツールが"
         "「テスト通過」を優先する設計によりバグ発見と矛盾することを指摘している。"
         "また、LLM生成テストにはAssertion RouletteやMagic Number Test等のテストスメルが頻出し [9]、"
         "テストアサーションの正確性は50%以下で「実装された動作」を捉えがちであることが報告されている [10]。")

    para(doc,
         "本研究はテスト生成の精度向上ではなく、テスト設計情報の構造化・外部化に焦点を当てる点で"
         "既存手法と目的が異なる。BDD/Specification by Example [11] は「非技術者にも可読なテスト仕様」"
         "という思想を共有するが、受入テスト・システムテスト層が対象であり、単体テスト層のコード構造に"
         "基づく品質情報の構造化は扱わない。ACL 2025の多段階テスト仕様生成 [12] は自然言語要求からの"
         "トップダウン方式であり、本研究の既存コードからのボトムアップ方式とは方向が異なる。")

    # ------------------------------------------------------------------
    # 2. 実施概要
    # ------------------------------------------------------------------
    heading(doc, "2. 実施概要")

    # 2.1 手法
    subheading(doc, "2.1 手法")

    para(doc,
         "提案手法は3段階のパイプラインである（図1）。各段階において大規模言語モデルを活用する。")

    bold_para(doc, [
        ("Step 1. 対象選定（静的解析 + LLM推論）: ", True),
        ("対象リポジトリのソースコード構造を解析し、テスト対象関数を選定する。選定基準は "
         "(a) 外部依存なし（純粋関数）、(b) 分岐構造が明確、(c) 既存テストが存在（比較用）の3点とした。"
         "LLMへの入力はリポジトリのディレクトリ構成とソースファイル一覧であり、出力は選定基準を満たす関数のリストである。", False),
    ])

    bold_para(doc, [
        ("Step 2. TRM生成（静的解析 + LLM推論）: ", True),
        ("選定した各関数のソースコードから分岐条件を抽出し、以下の5種別でテスト要求を定義する（表1）。"
         "LLMへの入力は対象関数のソースコード全文（平均48行）とTRM生成プロンプト"
         "（5種別の定義と出力スキーマを含む）であり、出力はYAML形式のTRM（平均12.4件/関数）である。", False),
    ])

    # 表1: TRM種別定義
    para(doc, "表1: TRM種別定義", bold=True, sz=9, align=WD_ALIGN_PARAGRAPH.CENTER, after=Pt(2))
    add_table(doc,
              ["種別", "記号", "定義"],
              [
                  ["分岐網羅", "BR", "if/else/switchの各パスを検証"],
                  ["同値クラス", "EC", "入力パラメータの同値分割を検証"],
                  ["境界値", "BV", "同値クラス境界の上下限を検証"],
                  ["エラーパス", "ER", "異常系・エラー処理を検証"],
                  ["依存切替", "DP", "関数間の依存関係・差異を検証"],
              ],
              col_widths=[2.5, 1.5, 8.0], fs=9)

    para(doc,
         "各要求には一意のID、自然言語の説明、優先度、根拠となるソースコード行番号を付与し、"
         "YAML形式で出力する。以下はParseVersion関数に対するTRMの抜粋である。")

    code_block(doc,
               '- id: "BR-11"\n'
               '  type: branch_coverage\n'
               '  description: "alpha修飾子の分岐を通過する"\n'
               '  priority: high\n'
               '  source_ref: "BC-02-10: *p == \'a\' && strncmp == \'alpha\'"\n'
               '  expected_behavior: "nShift = -0x60 を設定"\n'
               '\n'
               '- id: "ER-01"\n'
               '  type: error_path\n'
               '  description: "空文字列入力時にクラッシュしないことを検証する"\n'
               '  priority: high\n'
               '  source_ref: "関数先頭のポインタ参照"\n'
               '  expected_behavior: "デフォルト値 0x80808080 を返却"')

    bold_para(doc, [
        ("Step 3. テストコード生成（LLM）: ", True),
        ("TRMのYAML全体と対象ソースコードをLLMに入力し、Google Test準拠のC++テストコードを出力させる。"
         "各テストケースにはTRMのIDをコメントで記載し、要求からテストへの追跡性を確保する。", False),
    ])

    code_block(doc,
               'TEST(ParseVersion, AlphaModifier) {\n'
               '    // BR-11: alpha修飾子の分岐を通過する\n'
               '    UINT32 val = ParseVersion(L"2.4.1alpha");\n'
               '    EXPECT_EQ(val, static_cast<UINT32>(0x82848120));\n'
               '}')

    # 図1挿入（2.1節「図1」の言及直後）
    figure(doc, PNGS["fig1"],
           "図1: 提案手法の全体構成（3段階パイプライン）", Cm(14.5))

    # 2.2 実験条件
    subheading(doc, "2.2 実験条件")

    para(doc, "表2: 実験条件と対象関数", bold=True, sz=9, align=WD_ALIGN_PARAGRAPH.CENTER, after=Pt(2))
    add_table(doc,
              ["領域", "対象関数", "コード行数", "既存テスト数"],
              [
                  ["日時書式・バージョン解析\n(format.cpp)",
                   "GetDateTimeFormat,\nParseVersion,\nCompareVersion",
                   "約115行", "約20件"],
                  ["メールアドレス・文字種判定\n(CWordParse.cpp)",
                   "IsMailAddress,\nWhatKindOfTwoChars,\nWhatKindOfTwoChars4KW",
                   "約180行", "約15件"],
                  ["全角半角英数変換\n(convert_util.cpp)",
                   "Convert_ZeneisuToHaneisu,\nConvert_HaneisuToZeneisu",
                   "約90行", "約10件"],
                  ["合計", "8関数", "約385行", "約45件"],
              ],
              col_widths=[3.5, 4.0, 2.0, 2.0], fs=8.5)

    para(doc,
         "対象プロジェクト: sakura-editor/sakura（C++、OSSテキストエディタ、Google Test整備済み）")
    para(doc,
         "実行環境: macOS (Apple clang 16.0) + Google Test 1.17 + Windows互換ヘッダ"
         "（純粋関数のため型定義の互換レイヤーのみで移植）")
    para(doc, "使用モデル: 大規模言語モデル（匿名査読のためモデル名は伏せる）")
    para(doc, "比較対象: (A) 既存の手動テスト（約45件）、(C) 提案手法によるTRM + 生成テスト")

    # ------------------------------------------------------------------
    # 3. 実施結果
    # ------------------------------------------------------------------
    heading(doc, "3. 実施結果")

    # 3.1 TRM生成結果
    subheading(doc, "3.1 TRM生成結果（P1: テスト設計根拠の明示に対応）")

    para(doc,
         "8関数に対し計99件のテスト要求を定義した。種別ごとの内訳を表3に示す（図2）。")

    para(doc, "表3: TRM種別別内訳", bold=True, sz=9, align=WD_ALIGN_PARAGRAPH.CENTER, after=Pt(2))
    add_table(doc,
              ["種別", "件数", "割合", "要求あたりのコード行数"],
              [
                  ["BR（分岐網羅）", "55", "55.6%", "7.0行/要求"],
                  ["EC（同値クラス）", "27", "27.3%", "14.3行/要求"],
                  ["BV（境界値）", "11", "11.1%", "35.0行/要求"],
                  ["ER（エラーパス）", "3", "3.0%", "128.3行/要求"],
                  ["DP（依存切替）", "3", "3.0%", "128.3行/要求"],
                  ["合計", "99", "100%", "3.9行/要求"],
              ],
              col_widths=[3.5, 1.5, 1.5, 3.5], fs=9)

    para(doc,
         "BR要求が最多であるのは、ソースコードのif/else/switch分岐から直接導出されるためである。"
         "ER/DPが少ないのは対象が純粋関数であり、エラーハンドリングや関数間依存が限定的であることによる。"
         "領域別の圧縮率は format.cpp（2.5行/要求）、CWordParse.cpp（5.3行/要求）、"
         "convert_util.cpp（4.7行/要求）であり、分岐密度の高い関数ほどテスト要求密度が高く、"
         "TRMがコードの複雑度に比例した適切な粒度で品質要求を抽出していることを示す。")

    para(doc,
         "99件のテスト要求のdescription（自然言語記述）を3段階"
         "（L1: コード知識不要、L2: ドメイン知識のみ必要、L3: コード知識が必要）で分類した結果、"
         "L1+L2が65件（65.7%）であった。すなわち、テスト要求の約2/3はコードを読めない人でも"
         "内容を理解できる。種別間ではEC（81.5%）とER（100%）の可読率が高く、入力パターンの分類や"
         "エラー処理はドメインレベルで自然に表現できる概念であることを示している。"
         "一方、DP（33.3%）は関数間の関係性を記述するためコード知識が必要になりやすい。"
         "領域間ではformat.cpp（78.3%）が最も高く、CWordParse.cpp（41.2%）が最も低い。"
         "可読率はドメインの一般性と相関する。")

    # 図2挿入（3.1節「図2」の言及直後）
    figure(doc, PNGS["fig2"],
           "図2: テスト要求の種別内訳（N=99）", Cm(10))

    # 3.2 テスト生成結果
    subheading(doc, "3.2 テスト生成結果（P1に対応）")

    para(doc,
         "TRMを入力として計145件のGoogle Testテストコードを生成した。"
         "テスト要求99件に対するカバー率は100%（全要求に対応するテストケースが存在）であった。"
         "テストケース数が要求数を上回るのは、1つの境界値要求を上限・下限の別テストで検証する等のためである。"
         "既存の手動テスト約45件に対し約3.2倍のテストケースを生成した（図3）。")

    # 図3挿入（3.2節「図3」の言及直後）
    figure(doc, PNGS["fig3"],
           "図3: 既存テストと生成テストの領域別比較", Cm(12))

    # 3.3 コンパイル・実行結果
    subheading(doc, "3.3 コンパイル・実行結果")

    para(doc,
         "全3ファイルのコンパイル成功後、テスト実行を行った結果を以下に示す（図4）。")

    add_table(doc,
              ["指標", "値"],
              [
                  ["初回テストPASS", "133/145件 (91.7%)"],
                  ["初回テストFAIL", "12件 (8.3%)"],
                  ["FAIL原因: 対象関数のバグ", "0件"],
                  ["FAIL原因: LLMの期待値推論誤り", "12件 (100%)"],
                  ["修正後テストPASS", "145/145件 (100%)"],
              ],
              col_widths=[5.0, 5.0], fs=9)

    para(doc,
         "12件のFAILは全てLLMがテスト期待値を誤って推論したものであり、"
         "対象関数の実装バグは検出されなかった。")

    # 図4挿入（3.3節「図4」の言及直後）
    figure(doc, PNGS["fig4"],
           "図4: コンパイル・実行結果（初回→修正後）", Cm(12))

    # 3.4 FAIL分析とTRMの診断効果
    subheading(doc, "3.4 FAIL分析とTRMの診断効果（P3: テスト失敗時の原因特定コストに対応）")

    para(doc,
         "12件のFAILについて、TRMのIDから原因を分類した結果を以下に示す。"
         "TRM IDはmacos-build-reportにおける失敗テスト名との照合に基づく。")

    add_table(doc,
              ["原因分類", "件数", "対応するTRM ID", "原因の技術的詳細"],
              [
                  ["コンポーネント分割の誤解", "5", "BR-11〜15\n（BR-20〜23と関連）",
                   '"2.4.1alpha"を1コンポーネントと予測。\n実際は"1"と"alpha"が別コンポーネント'],
                  ["数字グルーピングの誤解", "1", "EC（同値クラス内部）",
                   '"1234"を1桁ずつと予測。\n実際は2桁ずつ区切り'],
                  ["文字数カウントミス", "2", "BR-34, EC\n（ドメイン構造）",
                   "wcslen=18をnull込み19と誤算"],
                  ["API呼び出しバグ", "1", "EC-16",
                   "offsetとバッファポインタの二重シフト"],
                  ["マッピング戻り値の誤解", "3", "BR-39〜41",
                   "変数マッピング後の返却値を元の値と誤認"],
              ],
              col_widths=[3.0, 1.0, 3.0, 5.0], fs=8.5)

    para(doc,
         "TRMがない場合、145件のテストコードとC++ソースコードを逐一突合し、"
         "16進数のビット操作やポインタ演算を追跡して原因を特定する必要がある。"
         "TRMがある場合、テストケースに付与されたTRM IDから該当要求の日本語説明と"
         "根拠（ソースコード行参照）を即座に特定でき、原因分類が構造的に行えた。"
         "12件全てが要求レベルで分類可能であり、うち約8件はコード非専門者にも"
         "理解可能な粒度で診断できた。")

    # 3.5 品質ギャップ分析
    subheading(doc, "3.5 品質ギャップ分析（P2: 網羅性の事前把握に対応）")

    para(doc,
         "TRMと既存テストを照合し、既存テストで未カバーのテスト要求を種別ごとに集計した（図5）。"
         "なお、既存テストとの照合は既存テストコードの内容からの推定に基づく。")

    add_table(doc,
              ["種別", "TRM要求数", "既存テストでカバー済み（推定）", "未カバー（推定）", "未カバー率"],
              [
                  ["BR", "55", "26", "29", "52.7%"],
                  ["EC", "27", "8", "19", "70.4%"],
                  ["BV", "11", "3", "8", "72.7%"],
                  ["ER", "3", "1", "2", "66.7%"],
                  ["DP", "3", "0", "3", "100%"],
                  ["合計", "99", "38", "61", "61.6%"],
              ],
              col_widths=[1.8, 2.0, 3.5, 2.5, 2.0], fs=8.5)

    para(doc,
         "種別間の差は顕著であり、DP(100%) > BV(72.7%) > EC(70.4%) > ER(66.7%) > BR(52.7%)"
         "の順で未カバー率が高い。DP（関数間の差異検証・往復変換検証）は既存テストに完全に欠落していた。"
         "BV（境界値）とEC（同値クラス）の未カバー率も70%を超えており、既存の手動テストが分岐の主要パス"
         "（BR）に偏っていることが定量的に示された。")

    para(doc,
         "この分析自体が「TRMが存在するからこそ可能」であることに注意すべきである。"
         "TRMがなければ、既存テストの品質ギャップは「テストが少ない」という漠然とした認識にとどまり、"
         "「何の観点が欠けているか」を構造的に把握することはできない。"
         "これはP2（網羅性の事前把握）への解決策として、TRMがテスト「前」に品質ギャップを"
         "種別ごとに特定できることを示す。")

    # 図5挿入（3.5節「図5」の言及直後）
    figure(doc, PNGS["fig5"],
           "図5: 既存テストの品質ギャップ分析（種別別未カバー率）", Cm(10))

    # 3.6 可読性分析
    subheading(doc, "3.6 可読性分析")

    para(doc,
         "TRMの可読性（コード非専門者がテスト要求を理解できるか）を種別別・領域別に"
         "分析した結果を以下に示す。")

    add_table(doc,
              ["種別", "件数", "L1（コード知識不要）", "L2（ドメイン知識のみ）",
               "L3（コード知識必要）", "可読率(L1+L2)"],
              [
                  ["BR", "55", "20", "14", "21", "61.8%"],
                  ["EC", "27", "13", "9", "5", "81.5%"],
                  ["BV", "11", "3", "5", "3", "72.7%"],
                  ["ER", "3", "3", "0", "0", "100.0%"],
                  ["DP", "3", "0", "1", "2", "33.3%"],
                  ["全体", "99", "38", "27", "34", "65.7%"],
              ],
              col_widths=[1.5, 1.0, 2.5, 2.5, 2.5, 2.0], fs=8)

    add_table(doc,
              ["領域", "件数", "可読率(L1+L2)"],
              [
                  ["format.cpp（日時書式・バージョン解析）", "46", "78.3%"],
                  ["convert_util.cpp（全角半角英数変換）", "19", "73.7%"],
                  ["CWordParse.cpp（メールアドレス・文字種判定）", "34", "41.2%"],
              ],
              col_widths=[6.0, 2.0, 3.0], fs=8.5)

    para(doc,
         "全体の約2/3（65.7%）がコード非専門者にも可読であった。"
         "一般に理解しやすいドメイン（日時書式、全角半角変換）では可読率が高く、"
         "エディタ固有のロジック（文字種結合判定のECharKind列挙型やマッピング規則等）を扱う関数では"
         "可読率が低い。なお、L3に分類された34件のうち一部は、description文の改善"
         "（例: 「ラテン→CSYMマッピング」を「ラテン系文字がアルファベットと同種として扱われること」に"
         "言い換え）によりL2に引き上げ可能と考えられる。")

    para(doc,
         "本分析は著者による分類であり、実際の非開発者による理解度評価は"
         "3.6節の制約として留意する必要がある。")

    # ------------------------------------------------------------------
    # 4. 結論
    # ------------------------------------------------------------------
    heading(doc, "4. 結論")

    # 4.1 まとめ
    subheading(doc, "4.1 まとめ")

    para(doc,
         "本研究では、ソースコードの静的解析からテスト要求モデル（TRM）をYAML形式で生成し、"
         "TRMを入力として大規模言語モデルに単体テストを生成させる手法を提案・実証した。"
         "OSSのC++コード8関数（約385行）に適用した結果:")

    bold_para(doc, [
        ("1. TRM生成: ", True),
        ("5種別99件のテスト要求を構造化定義。分岐密度に応じた適切な粒度（平均3.9行/要求）で"
         "品質要求を抽出し、そのうち65.7%がコード非専門者にも可読であった。", False),
    ])

    bold_para(doc, [
        ("2. テスト生成: ", True),
        ("TRMからの145件のテストが全要求をカバー。既存テストと比較して約3.2倍のテストケースを"
         "体系的に生成した。ただし、既存テストは本手法と異なる目的で作成されたものであり、"
         "件数の直接比較には留意が必要である。", False),
    ])

    bold_para(doc, [
        ("3. 実行精度: ", True),
        ("初回PASS率91.7%。FAIL 12件は全てLLMの期待値推論誤り（対象関数バグ0件）。"
         "TRM IDにより原因を5分類に即時特定し、修正後100% PASS。", False),
    ])

    bold_para(doc, [
        ("4. 品質ギャップ発見: ", True),
        ("既存テストの未カバー要求を推定61件（61.6%）、種別ごとに特定した。"
         "特にDP種別は100%未カバーであり、TRMによる構造化された品質ギャップの可視化が有効に機能した。", False),
    ])

    para(doc,
         "TRMの導入により、テスト設計根拠の明示（P1）、網羅性の事前定量管理（P2）、"
         "テスト失敗時の構造的原因分類（P3）を実現した。")

    # 4.2 今後の課題
    subheading(doc, "4.2 今後の課題")

    para(doc,
         "本実験は純粋関数8関数に限定されており、以下の課題が残る。")

    bold_para(doc, [
        ("(1) 副作用関数への拡張: ", True),
        ("本実験の対象は純粋関数に限定されており、ER/DPが計6件と少ない。"
         "モック・スタブを要する関数でのTRM生成ではER/DP比率の増加が予測され、"
         "TRMスキーマの拡張（モック条件の記述等）が必要である。", False),
    ])

    bold_para(doc, [
        ("(2) コードカバレッジとの相関分析: ", True),
        ("TRM要求カバー率100%が実際のC0/C1カバレッジとどう対応するかは未検証であり、"
         "両指標の相関分析によりTRMの品質保証能力の定量的裏付けが必要である。", False),
    ])

    bold_para(doc, [
        ("(3) 可読性の被験者評価: ", True),
        ("本研究の可読率65.7%は著者による分類であり、実際の非開発者"
         "（QAエンジニア、プロジェクトマネージャー等）による理解度評価実験が必要である。", False),
    ])

    bold_para(doc, [
        ("(4) 他言語・他プロジェクトでの再現実験: ", True),
        ("本実験はC++の1プロジェクト（sakura-editor）に限定されている。"
         "TRMの5種別体系は言語非依存に設計されているが、Java/Python等の他言語や"
         "Webアプリケーション等の異なるドメインでの再現実験により一般化可能性を検証する必要がある。", False),
    ])

    # ------------------------------------------------------------------
    # 参考文献
    # ------------------------------------------------------------------
    para(doc, "", sz=4, after=Pt(2))
    heading(doc, "参考文献")

    refs = [
        '[1] C. Lemieux, J. P. Inala, S. K. Lahiri, and S. Sen, '
        '"CodaMOSA: Escaping Coverage Plateaus in Test Generation with Pre-trained Large Language Models," '
        'in Proc. ICSE 2023, pp. 919-931, 2023.\n'
        'https://dl.acm.org/doi/10.1109/icse48619.2023.00085',

        '[2] Y. Chen et al., '
        '"ChatUniTest: A Framework for LLM-Based Test Generation," '
        'in Companion Proc. FSE 2024, ACM, 2024.\n'
        'https://dl.acm.org/doi/abs/10.1145/3663529.3663801',

        '[3] M. Schafer et al., '
        '"An Empirical Evaluation of Using Large Language Models for Automated Unit Test Generation," '
        'IEEE Transactions on Software Engineering, vol. 50, 2024.\n'
        'https://ieeexplore.ieee.org/document/10329992/',

        '[4] J. A. Ramos, E. M. Bader, J. Bader, and S. Iyer, '
        '"CoverUp: Coverage-Guided LLM-Based Test Generation," '
        'arXiv:2403.16218, 2024.\n'
        'https://arxiv.org/abs/2403.16218',

        '[5] Qodo (CodiumAI), '
        '"Qodo-Cover: An AI-Powered Tool for Automated Test Generation and Code Coverage Enhancement," '
        'GitHub, 2024.\n'
        'https://github.com/Codium-ai/cover-agent',

        '[6] "Enhancing LLM-based Test Generation for Hard-to-Cover Branches via Program Analysis (TELPA)," '
        'arXiv:2404.04966, ACM TOSEM, 2024.\n'
        'https://dl.acm.org/doi/10.1145/3748505',

        '[7] B. Chu, Y. Feng, K. Liu, Z. Guo, Y. Zhang, H. Shi, Z. Nan, and B. Xu, '
        '"Large Language Models for Unit Test Generation: Achievements, Challenges, and the Road Ahead," '
        'arXiv:2511.21382, 2025.\n'
        'https://arxiv.org/abs/2511.21382',

        '[8] N. S. Mathews and M. Nagappan, '
        '"Design choices made by LLM-based test generators prevent them from finding bugs," '
        'arXiv:2412.14137, 2024.\n'
        'https://arxiv.org/abs/2412.14137',

        '[9] "Test smells in LLM-Generated Unit Tests," '
        'arXiv:2410.10628, 2024.\n'
        'https://arxiv.org/abs/2410.10628',

        '[10] "Do LLMs generate test oracles that capture the actual or the expected program behaviour?" '
        'arXiv:2410.21136, 2024.\n'
        'https://arxiv.org/abs/2410.21136',

        '[11] G. Adzic, Specification by Example: How Successful Teams Deliver the Right Software, '
        'Manning, 2011.',

        '[12] "Multi-Step Generation of Test Specifications using Large Language Models," '
        'in Proc. ACL 2025 Industry Track, 2025.\n'
        'https://aclanthology.org/2025.acl-industry.11.pdf',

        '[13] C. Wiecher et al., '
        '"Model-based analysis and specification of functional requirements and tests '
        'for complex automotive systems," '
        'Systems Engineering, Wiley, 2024.\n'
        'https://incose.onlinelibrary.wiley.com/doi/full/10.1002/sys.21748',

        '[14] "Static Analysis as a Feedback Loop: '
        'Enhancing LLM-Generated Code Beyond Correctness," '
        'arXiv:2508.14419, 2025.\n'
        'https://arxiv.org/abs/2508.14419',

        '[15] 「生成AIを活用したテストパターンマトリックスを用いたテスト生成」, '
        'SQiP 2024 一般発表 A4-1, 2024.\n'
        'https://www.juse.jp/sqip/symposium/2024/timetable/files/A4-1_ronbun.pdf',

        '[16] 産業技術総合研究所 (AIST), '
        '「生成AI品質マネジメントガイドライン第1版」, Rev. 1.0.0, 2025年5月.\n'
        'https://www.digiarc.aist.go.jp/publication/aiqm/genaiqm-guidelines-v1.html',

        '[17] IPA 独立行政法人情報処理推進機構, '
        '「テキスト生成AIの導入・運用ガイドライン」, 2024年7月.\n'
        'https://www.ipa.go.jp/jinzai/ics/core_human_resource/final_project/2024/generative-ai-guideline.html',
    ]

    for ref in refs:
        para(doc, ref, sz=8, after=Pt(3))

    # ========== 保存 ==========
    doc.save(OUT)
    print(f"\nSaved: {OUT}")
    print(f"File size: {os.path.getsize(OUT):,} bytes")

    # ========== 埋め込み画像サイズの検証 ==========
    print("\n--- 埋め込み画像サイズの検証 ---")
    with zipfile.ZipFile(OUT, 'r') as z:
        image_files = [n for n in z.namelist() if n.startswith('word/media/')]
        assert len(image_files) == 5, f"Expected 5 embedded images, found {len(image_files)}"
        for img_name in sorted(image_files):
            info = z.getinfo(img_name)
            size_bytes = info.file_size
            print(f"  {img_name}: {size_bytes:,} bytes")
            assert size_bytes >= 70000, (
                f"Embedded image too small (<70KB): {img_name} ({size_bytes:,} bytes)"
            )
    print("\nAll 5 embedded images are >= 70KB. Verification passed.")


if __name__ == "__main__":
    main()
