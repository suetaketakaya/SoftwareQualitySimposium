#!/usr/bin/env python3
"""SQiP 2026 最終提出用 docx: abstract-final-v2.md を忠実にdocx化（cairosvg不使用）"""
import os
import zipfile
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE = "/Users/suetaketakaya/SoftwareQualitySymposium"
FIG = os.path.join(BASE, "drafts", "figures")
OUT = os.path.join(BASE, "report", "submission_final_v2_2026.docx")

# matplotlib生成済みPNGのパスを直接参照（cairosvg不使用・SVG変換一切なし）
PNGS = {
    "fig1": os.path.join(FIG, "fig1-overview.png"),
    "fig2": os.path.join(FIG, "fig2-requirements-breakdown.png"),
    "fig3": os.path.join(FIG, "fig3-test-comparison.png"),
    "fig4": os.path.join(FIG, "fig4-execution-results.png"),
    "fig5": os.path.join(FIG, "fig5-quality-gap.png"),
}

# ========== PNGファイルの存在確認・サイズ検証（70KB以上） ==========
for k, path in PNGS.items():
    assert os.path.exists(path), f"PNG not found: {path}"
    size = os.path.getsize(path)
    assert size >= 70000, f"PNG too small (<70KB): {path} ({size} bytes)"
    print(f"  OK: {k} = {size:,} bytes ({path})")


# ========== ヘルパー関数 ==========

def set_font(run, name="游明朝", size=10, bold=False, italic=False, color=None):
    """フォント設定の共通ヘルパー"""
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    # 日本語フォントの設定
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = run._element.makeelement(qn('w:rFonts'), {})
        rpr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    if color:
        run.font.color.rgb = color


def heading(doc, text, lv=2):
    """セクション見出し"""
    x = doc.add_heading(text, level=lv)
    for r in x.runs:
        set_font(r, size=13 if lv == 1 else (11 if lv == 2 else 10.5),
                 bold=True, color=RGBColor(0, 0, 0))
    x.paragraph_format.space_before = Pt(12 if lv <= 2 else 8)
    x.paragraph_format.space_after = Pt(4)
    x.paragraph_format.line_spacing = Pt(16)
    return x


def subheading(doc, text):
    """サブセクション見出し（### レベル）"""
    x = doc.add_paragraph()
    r = x.add_run(text)
    set_font(r, size=10.5, bold=True)
    x.paragraph_format.space_before = Pt(8)
    x.paragraph_format.space_after = Pt(2)
    x.paragraph_format.line_spacing = Pt(16)
    return x


def para(doc, text, sz=10, after=Pt(4), bold=False, align=None, italic=False):
    """本文段落"""
    x = doc.add_paragraph()
    r = x.add_run(text)
    set_font(r, size=sz, bold=bold, italic=italic)
    if align:
        x.alignment = align
    x.paragraph_format.space_after = after
    x.paragraph_format.space_before = Pt(0)
    x.paragraph_format.line_spacing = Pt(16)
    return x


def bold_para(doc, text_parts, sz=10, after=Pt(4)):
    """太字と通常テキストの混在段落。text_partsは(text, bold)のリスト"""
    x = doc.add_paragraph()
    for txt, is_bold in text_parts:
        r = x.add_run(txt)
        set_font(r, size=sz, bold=is_bold)
    x.paragraph_format.space_after = after
    x.paragraph_format.space_before = Pt(0)
    x.paragraph_format.line_spacing = Pt(16)
    return x


def italic_bold_para(doc, text_parts, sz=10, after=Pt(4)):
    """太字/イタリック/通常テキストの混在段落。text_partsは(text, bold, italic)のリスト"""
    x = doc.add_paragraph()
    for item in text_parts:
        if len(item) == 3:
            txt, is_bold, is_italic = item
        else:
            txt, is_bold = item
            is_italic = False
        r = x.add_run(txt)
        set_font(r, size=sz, bold=is_bold, italic=is_italic)
    x.paragraph_format.space_after = after
    x.paragraph_format.space_before = Pt(0)
    x.paragraph_format.line_spacing = Pt(16)
    return x


def code_block(doc, text):
    """コードブロック（Consolas 8.5pt）"""
    x = doc.add_paragraph()
    r = x.add_run(text)
    r.font.size = Pt(8.5)
    r.font.name = "Consolas"
    x.paragraph_format.space_after = Pt(4)
    x.paragraph_format.space_before = Pt(2)
    x.paragraph_format.left_indent = Cm(0.8)
    x.paragraph_format.line_spacing = Pt(13)
    # 背景色（薄いグレー）をシェーディングで設定
    shading = x._element.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): 'F0F0F0'
    })
    x._element.get_or_add_pPr().append(shading)


def add_table(doc, headers, rows, col_widths=None, fs=8.5):
    """表の追加"""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # ヘッダー行
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        set_font(r, size=fs, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].paragraph_format.space_after = Pt(1)
        cell.paragraphs[0].paragraph_format.space_before = Pt(1)
        # ヘッダー背景色
        shading = cell._element.makeelement(qn('w:shd'), {
            qn('w:val'): 'clear',
            qn('w:color'): 'auto',
            qn('w:fill'): 'D9E2F3'
        })
        cell._element.get_or_add_tcPr().append(shading)
    # データ行
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = ""
            r = cell.paragraphs[0].add_run(str(val))
            set_font(r, size=fs)
            cell.paragraphs[0].paragraph_format.space_after = Pt(1)
            cell.paragraphs[0].paragraph_format.space_before = Pt(1)
    if col_widths:
        for ri, row in enumerate(t.rows):
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Cm(w)
    return t


def add_figure(doc, png_key, caption, width=Cm(14)):
    """図を追加（中央揃え + キャプション）"""
    # 画像
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(PNGS[png_key], width=width)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    # キャプション
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    set_font(r, size=9)
    cap.paragraph_format.space_after = Pt(6)
    cap.paragraph_format.space_before = Pt(0)
    cap.paragraph_format.line_spacing = Pt(14)


def add_page_break(doc):
    """改ページ"""
    p = doc.add_paragraph()
    run = p.add_run()
    run._element.makeelement(qn('w:br'), {qn('w:type'): 'page'})
    br = run._element.makeelement(qn('w:br'), {qn('w:type'): 'page'})
    run._element.append(br)


# ========== ドキュメント生成 ==========
doc = Document()

# デフォルトスタイルの設定
style = doc.styles['Normal']
font = style.font
font.name = '游明朝'
font.size = Pt(10)
style.paragraph_format.line_spacing = Pt(16)
# 日本語フォント設定
rpr = style.element.get_or_add_rPr()
rFonts = rpr.find(qn('w:rFonts'))
if rFonts is None:
    rFonts = style.element.makeelement(qn('w:rFonts'), {})
    rpr.insert(0, rFonts)
rFonts.set(qn('w:eastAsia'), '游明朝')

# ページ余白の設定
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ================================================================
# ページ1: 発表申込書
# ================================================================
para(doc, "ソフトウェア品質シンポジウム 2026　発表申込書", sz=12, bold=True,
     align=WD_ALIGN_PARAGRAPH.CENTER, after=Pt(16))

# タイトル
para(doc, "■ 発表タイトル", sz=10, bold=True, after=Pt(2))
para(doc, "静的解析に基づくテスト要求モデルの自動生成と、それを用いたLLMテスト生成の実証評価",
     sz=10, after=Pt(10))

# 申込区分
para(doc, "■ 申込区分", sz=10, bold=True, after=Pt(2))
para(doc, "経験発表", sz=10, after=Pt(10))

# カテゴリ
para(doc, "■ カテゴリ", sz=10, bold=True, after=Pt(2))
para(doc, "品質管理・テスト技術の観点", sz=10, after=Pt(10))

# キーワード
para(doc, "■ キーワード", sz=10, bold=True, after=Pt(2))
para(doc, "ホワイトボックステスト / テスト要求モデル / LLM / 静的解析 / 単体テスト生成 / トレーサビリティ",
     sz=10, after=Pt(10))

# 改ページ
add_page_break(doc)

# ================================================================
# ページ2以降: アブストラクト全文
# ================================================================

# タイトル（中央揃え）
para(doc, "静的解析に基づくテスト要求モデルの自動生成と、それを用いたLLMテスト生成の実証評価",
     sz=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=Pt(16))

# ================================================================
# 1. ねらい
# ================================================================
heading(doc, "1. ねらい", lv=2)

# 1.1 背景
subheading(doc, "1.1 背景")

para(doc, "大規模言語モデル（LLM）の急速な発展により、ソフトウェア開発のワークフローは大きく変化しつつある。"
     "現在のLLM活用開発の主流は、荒い要求仕様からLLMに実装コードを直接生成させるパターンであり、"
     "評価が「要求から現れる基本機能の動作確認」に偏る傾向がある。しかし、LLMが生成したコードを実リリースに"
     "載せるためには、単体テスト・コンポーネントテストの層で品質を担保する必要がある。")

para(doc, "現状のLLMテスト生成手法（CodaMOSA [1]、ChatUniTest [2]、TestPilot [3]、CoverUp [4]、"
     "CoverAgent [5]、TELPA [6] 等）はカバレッジ向上やコンパイル成功率の改善に注力している。"
     "Chu et al. [7] による115本の体系的サーベイでは、LLMテスト生成研究の89%が「プロンプトエンジニアリング"
     "による直接テスト生成」方式を採り、反復的な検証・修復ループで精度を改善するアーキテクチャが標準となっている。"
     "このアーキテクチャでは、テスト設計の意図はLLM内部に閉じたまま外部化されず、以下の3つの問題が構造的に未解決である。")

# 1.2 問題定義
subheading(doc, "1.2 問題定義")

bold_para(doc, [
    ("(P1) テスト設計根拠の不在", True),
    (": LLMが生成したテストが「何の分岐を」「どの同値クラスを」対象としているか不明であり、"
     "テスト設計の妥当性を第三者が検証できない。レビュー担当者は生成されたテストの妥当性を判断できず、"
     "結果として「LLMが作ったテストだから信頼できない」という属人的判断に陥る。", False)
])

bold_para(doc, [
    ("(P2) 網羅性の事前把握が不可能", True),
    (": テスト群がソースコードのどの範囲をカバーしているかは、コンパイル・実行してカバレッジツールを通すまで分からない。"
     "テスト生成「前」に網羅性を設計・管理する手段がなく、テスト件数が多くても同じパスの重複テストばかりという"
     "状況を事前に検知できない。", False)
])

bold_para(doc, [
    ("(P3) テスト失敗時の原因特定コスト", True),
    (": テスト失敗時に、生成テストのアサーションとソースコードを逐一突合する必要があり、"
     "原因分類に時間がかかる。例えば145件のテストが一斉に失敗した場合、テストコードとソースコードの突合に長時間を要する。", False)
])

# 1.3 目的と提案
subheading(doc, "1.3 目的と提案")

bold_para(doc, [
    ("本研究は、これら3つの問題に対し、", False),
    ("ソースコードの静的解析結果からホワイトボックステスト要求モデル（以下TRM: Test Requirement Model）を"
     "YAML形式で自動生成し、TRMを入力としてLLMにテストコードを生成させる手法", True),
    ("を提案する。TRMは分岐条件・同値クラス・境界値をID付きで構造化した中間成果物であり、"
     "テスト生成の「前」にテスト設計の網羅性を定量評価できる。さらに、TRMに対する独立した網羅性監査と、"
     "監査結果に基づくテスト設計書の生成を行い、C0/C1/MC-DCカバレッジの達成を検証する。"
     "OSSプロジェクトのC++コード8関数に適用し、TRMの生成精度、網羅性監査プロセスの有効性、"
     "テスト生成品質、テスト失敗時の診断効率を定量評価した。", False)
])

# 1.4 先行研究との差異
subheading(doc, "1.4 先行研究との差異")

para(doc, "先行手法はいずれも「LLMに直接テストコードを生成させ、その精度を反復的に改善する」アーキテクチャを採る [7]。"
     "Mathews & Nagappan [8] は、CoverAgent等のLLMテスト生成ツールが「テスト通過」を優先する設計により"
     "バグ発見と矛盾することを指摘している。また、LLM生成テストにはAssertion RouletteやMagic Number Test等の"
     "テストスメルが頻出し [9]、テストアサーションの正確性は50%以下で「実装された動作」を捉えがちであることが"
     "報告されている [10]。")

para(doc, "本研究はテスト生成の精度向上ではなく、テスト設計情報の構造化・外部化に焦点を当てる点で"
     "既存手法と目的が異なる。BDD/Specification by Example [11] は「非技術者にも可読なテスト仕様」という"
     "思想を共有するが、受入テスト・システムテスト層が対象であり、単体テスト層のコード構造に基づく品質情報の"
     "構造化は扱わない。ACL 2025の多段階テスト仕様生成 [12] は自然言語要求からのトップダウン方式であり、"
     "本研究の既存コードからのボトムアップ方式とは方向が異なる。")

# 1.5 本研究の新規性
subheading(doc, "1.5 本研究の新規性")

para(doc, "先行研究39件の体系的調査に基づき、本研究の新規性を以下の3点に整理する。")

bold_para(doc, [
    ("新規性1: テスト要求モデル（TRM）という構造化中間成果物の導入", True),
    ("。LLMテスト生成研究（ICSE, FSE, ICSME等の主要会議で発表された手法 [1]-[6]）はいずれもテストコードを"
     "直接生成し、テスト設計の意図を外部化する中間成果物を持たない。本研究はYAML形式のTRMを導入し、"
     "テスト設計の意図をLLM内部から外部化した初の手法である。TRMは5種別（BR: 分岐網羅、EC: 同値クラス、"
     "BV: 境界値、ER: エラーパス、DP: 依存切替）でテスト要求を分類し、各要求にIDを付与してテストコードへの"
     "トレーサビリティを維持する。", False)
])

bold_para(doc, [
    ("新規性2: LLM生成テスト設計に対する網羅性監査プロセスの提案と実証", True),
    ("。LLMが生成したテストの品質を評価する研究は多いが、LLMが生成した「テスト設計」の品質を監査する研究は"
     "存在しない。本研究はTRM網羅性監査を実施し、初回TRM（99件）の漏れ64件を特定して163件に拡充した"
     "（分岐網羅率87%→推定97%）。この監査プロセスにより、LLMによるテスト設計の限界を定量的に示すとともに、"
     "改善の方法論を提供する。", False)
])

bold_para(doc, [
    ("新規性3: テスト品質の種別別可視化", True),
    ("。既存テストの品質ギャップを5種別で定量化し（BR 52.7%、EC 70.4%、BV 72.7%、ER 66.7%、DP 100%が"
     "未カバー）、テスト品質の構造的な偏りを可視化した。この分析はTRMが存在するからこそ可能であり、"
     "従来のカバレッジ数値（C0/C1等）では把握できない「何の観点が欠けているか」を明らかにする。", False)
])

# ================================================================
# 2. 実施概要
# ================================================================
heading(doc, "2. 実施概要", lv=2)

# 2.1 手法
subheading(doc, "2.1 手法")

para(doc, "提案手法は5段階のパイプラインである（図1）。各段階において大規模言語モデルを活用する。")

# Step 1
bold_para(doc, [
    ("Step 1. 対象選定（静的解析 + LLM推論）", True),
    (": 対象リポジトリのソースコード構造を解析し、テスト対象関数を選定する。選定基準は "
     "(a) 外部依存なし（純粋関数）、(b) 分岐構造が明確、(c) 既存テストが存在（比較用）の3点とした。"
     "LLMへの入力はリポジトリのディレクトリ構成とソースファイル一覧であり、出力は選定基準を満たす関数のリストである。", False)
])

# Step 2
bold_para(doc, [
    ("Step 2. TRM生成（静的解析 + LLM推論）", True),
    (": 選定した各関数のソースコードから分岐条件を抽出し、以下の5種別でテスト要求を定義する（表1）。"
     "LLMへの入力は対象関数のソースコード全文（平均48行）とTRM生成プロンプト（5種別の定義と出力スキーマを含む）"
     "であり、出力はYAML形式のTRM（平均12.4件/関数）である。", False)
])

# 表1: TRM種別定義
para(doc, "表1: TRM種別定義", sz=9, bold=True, after=Pt(2))
add_table(doc,
          ["種別", "記号", "定義"],
          [
              ["分岐網羅", "BR", "if/else/switchの各パスを検証"],
              ["同値クラス", "EC", "入力パラメータの同値分割を検証"],
              ["境界値", "BV", "同値クラス境界の上下限を検証"],
              ["エラーパス", "ER", "異常系・エラー処理を検証"],
              ["依存切替", "DP", "関数間の依存関係・差異を検証"],
          ],
          col_widths=[3.0, 2.0, 10.0])

para(doc, "各要求には一意のID、自然言語の説明、優先度、根拠となるソースコード行番号を付与し、YAML形式で出力する。"
     "以下はParseVersion関数に対するTRMの抜粋である。", after=Pt(2))

# YAMLコードブロック
code_block(doc,
           '- id: "BR-11"\n'
           '  type: branch_coverage\n'
           '  description: "alpha修飾子の分岐を通過する"\n'
           '  priority: high\n'
           '  source_ref: "BC-02-10: *p == \'a\' && strncmp == \'alpha\'"\n'
           '  expected_behavior: "nShift = -0x60 を設定"\n'
           '\n'
           '- id: "ER-01"\n'
           '  type: error_path\n'
           '  description: "空文字列入力時にクラッシュしないことを検証する"\n'
           '  priority: high\n'
           '  source_ref: "関数先頭のポインタ参照"\n'
           '  expected_behavior: "デフォルト値 0x80808080 を返却"')

# Step 3
bold_para(doc, [
    ("Step 3. TRM網羅性監査（静的解析 + 独立レビュー）", True),
    (": 生成されたTRMに対し、ホワイトボックステスト網羅性の観点から独立した監査を実施する。"
     "ソースコードの全分岐条件をTRMのBR要求と1対1で照合し、漏れている分岐条件・同値クラス・境界値・"
     "エラーパス・依存切替を特定する。監査結果に基づき追加のテスト要求をTRMに統合する。"
     "この段階で関数ごとのテスト設計書（分岐条件一覧・TRM対応表・カバレッジ達成計画を含む正式なテスト設計文書）を生成する。", False)
])

# Step 4
bold_para(doc, [
    ("Step 4. テストコード生成（LLM）", True),
    (": 監査済みTRMのYAML全体と対象ソースコードをLLMに入力し、Google Test準拠のC++テストコードを出力させる。"
     "各テストケースにはTRMのIDをコメントで記載し、要求からテストへの追跡性を確保する。", False)
])

# C++コードブロック
code_block(doc,
           'TEST(ParseVersion, AlphaModifier) {\n'
           '    // BR-11: alpha修飾子の分岐を通過する\n'
           '    UINT32 val = ParseVersion(L"2.4.1alpha");\n'
           '    EXPECT_EQ(val, static_cast<UINT32>(0x82848120));\n'
           '}')

# Step 5
bold_para(doc, [
    ("Step 5. カバレッジ精密分析（静的解析）", True),
    (": 生成されたテストスイート全体に対し、C0（ステートメントカバレッジ）、C1（ブランチカバレッジ）、"
     "MC/DC（Modified Condition/Decision Coverage）、ループ境界の各基準でカバレッジを精密分析する。"
     "各分岐条件・複合条件に対してどのテストケースがカバーしているかを1対1で対応付け、未カバー項目を特定する。", False)
])

# 図1を配置（Step 5の後 = 手法全体の説明後）
add_figure(doc, "fig1", "図1: 提案手法の5段階パイプライン概要")

# 2.2 実験条件
subheading(doc, "2.2 実験条件")

para(doc, "表2: 実験条件と対象関数", sz=9, bold=True, after=Pt(2))
add_table(doc,
          ["領域", "対象関数", "コード行数", "既存テスト数"],
          [
              ["日時書式・バージョン解析 (format.cpp)", "GetDateTimeFormat, ParseVersion, CompareVersion", "約115行", "約20件"],
              ["メールアドレス・文字種判定 (CWordParse.cpp)", "IsMailAddress, WhatKindOfTwoChars, WhatKindOfTwoChars4KW", "約180行", "約15件"],
              ["全角半角英数変換 (convert_util.cpp)", "Convert_ZeneisuToHaneisu, Convert_HaneisuToZeneisu", "約90行", "約10件"],
              ["合計", "8関数", "約385行", "約45件"],
          ],
          col_widths=[5.5, 5.5, 2.5, 2.5])

# 箇条書き
para(doc, "・対象プロジェクト: sakura-editor/sakura（C++、OSSテキストエディタ、Google Test整備済み）")
para(doc, "・実行環境: macOS (Apple clang 16.0) + Google Test 1.17 + Windows互換ヘッダ（純粋関数のため型定義の互換レイヤーのみで移植）")
para(doc, "・使用モデル: 大規模言語モデル（匿名査読のためモデル名は伏せる）")
para(doc, "・比較対象: (A) 既存の手動テスト（約45件）、(C) 提案手法によるTRM + 生成テスト")

# ================================================================
# 3. 実施結果
# ================================================================
heading(doc, "3. 実施結果", lv=2)

# 3.1 TRM生成結果
subheading(doc, "3.1 TRM生成結果（P1: テスト設計根拠の明示に対応）")

para(doc, "8関数に対し計99件のテスト要求を初回生成した。種別ごとの内訳を表3に示す（図2）。")

para(doc, "表3: TRM種別別内訳（初回生成）", sz=9, bold=True, after=Pt(2))
add_table(doc,
          ["種別", "件数", "割合", "要求あたりのコード行数"],
          [
              ["BR（分岐網羅）", "55", "55.6%", "7.0行/要求"],
              ["EC（同値クラス）", "27", "27.3%", "14.3行/要求"],
              ["BV（境界値）", "11", "11.1%", "35.0行/要求"],
              ["ER（エラーパス）", "3", "3.0%", "128.3行/要求"],
              ["DP（依存切替）", "3", "3.0%", "128.3行/要求"],
              ["合計", "99", "100%", "3.9行/要求"],
          ],
          col_widths=[4.0, 2.5, 2.5, 5.0])

# 図2を配置（表3の直後）
add_figure(doc, "fig2", "図2: TRM種別別内訳")

para(doc, "BR要求が最多であるのは、ソースコードのif/else/switch分岐から直接導出されるためである。"
     "ER/DPが少ないのは対象が純粋関数であり、エラーハンドリングや関数間依存が限定的であることによる。"
     "領域別の圧縮率は format.cpp（2.5行/要求）、CWordParse.cpp（5.3行/要求）、convert_util.cpp（4.7行/要求）"
     "であり、分岐密度の高い関数ほどテスト要求密度が高く、TRMがコードの複雑度に比例した適切な粒度で品質要求を"
     "抽出していることを示す。")

para(doc, "99件のテスト要求のdescription（自然言語記述）を3段階（L1: コード知識不要、L2: ドメイン知識のみ必要、"
     "L3: コード知識が必要）で分類した結果、L1+L2が65件（65.7%）であった。すなわち、テスト要求の約2/3は"
     "コードを読めない人でも内容を理解できる。種別間ではEC（81.5%）とER（100%）の可読率が高く、入力パターンの"
     "分類やエラー処理はドメインレベルで自然に表現できる概念であることを示している。一方、DP（33.3%）は関数間の"
     "関係性を記述するためコード知識が必要になりやすい。領域間ではformat.cpp（78.3%）が最も高く、"
     "CWordParse.cpp（41.2%）が最も低い。可読率はドメインの一般性と相関する。")

# 3.2 TRM網羅性監査結果
subheading(doc, "3.2 TRM網羅性監査結果（新規性2に対応）")

para(doc, "初回生成TRM（99件）に対してホワイトボックステスト網羅性の観点から独立監査を実施した。"
     "ソースコードの全分岐条件（91件）とTRMのBR要求を1対1で照合した結果、初回の分岐網羅率は87%（79/91分岐）であった。")

para(doc, "表4: 関数ごとの初回分岐網羅率", sz=9, bold=True, after=Pt(2))
add_table(doc,
          ["関数", "ソース分岐数", "TRM BR対応数", "分岐網羅率"],
          [
              ["GetDateTimeFormat", "14", "14", "100%"],
              ["ParseVersion", "20", "17", "85%"],
              ["CompareVersion", "3", "3", "100%"],
              ["IsMailAddress", "17", "14", "82%"],
              ["WhatKindOfTwoChars", "15", "12", "80%"],
              ["WhatKindOfTwoChars4KW", "10", "8", "80%"],
              ["Convert_ZeneisuToHaneisu", "6", "6", "100%"],
              ["Convert_HaneisuToZeneisu", "6", "5", "83%"],
              ["合計", "91", "79", "87%"],
          ],
          col_widths=[5.0, 3.0, 3.0, 3.0])

para(doc, "監査により5カテゴリ合計64件の追加要求を特定した（表5）。")

para(doc, "表5: 監査による追加要求の内訳", sz=9, bold=True, after=Pt(2))
add_table(doc,
          ["カテゴリ", "初回", "追加", "監査後合計"],
          [
              ["BR（分岐網羅）", "55", "10", "65"],
              ["EC（同値クラス）", "27", "21", "48"],
              ["BV（境界値）", "11", "18", "29"],
              ["ER（エラーパス）", "3", "10", "13"],
              ["DP（依存切替）", "3", "5", "8"],
              ["合計", "99", "64", "163"],
          ],
          col_widths=[4.0, 3.0, 3.0, 3.0])

para(doc, "追加要求の重要度別内訳はCritical 4件、High 11件、Medium 36件、Low 13件であった。"
     "Critical指摘にはGetDateTimeFormatのstr[6]バッファオーバーフロー（wYear 6桁以上での切り捨て）、"
     "CompareVersionのUINT32差分のintキャスト時オーバーフロー（符号反転の可能性）、"
     "IsMailAddressのnBufLen小値時の符号問題（size_tラップアラウンド）が含まれる。")

para(doc, "監査後の推定分岐網羅率は97%に改善された。この結果は、LLMによるTRM生成が主要分岐パスの抽出に"
     "有効である一方、境界値（11→29件、2.6倍増）やエラーパス（3→13件、4.3倍増）の抽出に構造的な弱点を持つことを"
     "定量的に示している。")

# 3.3 テスト生成結果
subheading(doc, "3.3 テスト生成結果（P1に対応）")

# 図3を配置（テスト生成結果の冒頭）
add_figure(doc, "fig3", "図3: テスト生成の比較")

para(doc, "監査済みTRM（163件）を入力として計248件のGoogle Testテストコードを生成した。"
     "初回TRM（99件）に対応する145件と、追加要求（64件）に対応する103件で構成される。"
     "テスト要求163件に対するカバー率は100%（全要求に対応するテストケースが存在）であった。"
     "テストケース数が要求数を上回るのは、1つの境界値要求を上限・下限の別テストで検証する等のためである"
     "（展開倍率: 平均2.5倍）。既存の手動テスト約45件に対し約5.5倍のテストケースを体系的に生成した。"
     "ただし、既存テストは本手法と異なる目的で作成されたものであり、件数の直接比較には留意が必要である。")

para(doc, "この段階で、関数ごとのテスト設計書を正式なテスト設計文書として生成した。テスト設計書には"
     "ソースコードの全分岐条件一覧、TRM要求との対応表、各テストケースの入力・期待出力・カバーする分岐条件が"
     "記載されており、テスト設計のレビューとトレーサビリティの基盤となる。")

# 3.4 コンパイル・実行結果
subheading(doc, "3.4 コンパイル・実行結果")

para(doc, "全3ファイルのコンパイル成功後、テスト実行を行った結果を以下に示す（図4）。")

add_table(doc,
          ["指標", "値"],
          [
              ["初回テストPASS", "133/145件 (91.7%)"],
              ["初回テストFAIL", "12件 (8.3%)"],
              ["FAIL原因: 対象関数のバグ", "0件"],
              ["FAIL原因: LLMの期待値推論誤り", "12件 (100%)"],
              ["修正後テストPASS", "145/145件 (100%)"],
          ],
          col_widths=[6.0, 6.0])

# 図4を配置
add_figure(doc, "fig4", "図4: テスト実行結果")

para(doc, "12件のFAILは全てLLMがテスト期待値を誤って推論したものであり、対象関数の実装バグは検出されなかった。"
     "なお、初回PASS率91.7%は初回TRM（99件）から生成された145件に対する値である。"
     "追加要求から生成されたテストを含む全248件の最終PASS率は100%である。")

# 3.5 FAIL分析とTRMの診断効果
subheading(doc, "3.5 FAIL分析とTRMの診断効果（P3: テスト失敗時の原因特定コストに対応）")

para(doc, "12件のFAILについて、TRMのIDから原因を分類した結果を以下に示す。")

add_table(doc,
          ["原因分類", "件数", "対応するTRM ID", "原因の技術的詳細"],
          [
              ["コンポーネント分割の誤解", "5", "BR-11〜15（BR-20〜23と関連）",
               '"2.4.1alpha"を1コンポーネントと予測。実際は"1"と"alpha"が別コンポーネント'],
              ["数字グルーピングの誤解", "1", "EC（同値クラス内部）",
               '"1234"を1桁ずつと予測。実際は2桁ずつ区切り'],
              ["文字数カウントミス", "2", "BR-34, EC（ドメイン構造）",
               "wcslen=18をnull込み19と誤算"],
              ["API呼び出しバグ", "1", "EC-16",
               "offsetとバッファポインタの二重シフト"],
              ["マッピング戻り値の誤解", "3", "BR-39〜41",
               "変数マッピング後の返却値を元の値と誤認"],
          ],
          col_widths=[3.5, 1.5, 4.0, 5.5],
          fs=8)

para(doc, "TRMがない場合、145件のテストコードとC++ソースコードを逐一突合し、16進数のビット操作やポインタ演算を"
     "追跡して原因を特定する必要がある。TRMがある場合、テストケースに付与されたTRM IDから該当要求の日本語説明と"
     "根拠（ソースコード行参照）を即座に特定でき、原因分類が構造的に行えた。12件全てが要求レベルで分類可能であり、"
     "うち約8件はコード非専門者にも理解可能な粒度で診断できた。")

# 3.6 品質ギャップ分析
subheading(doc, "3.6 品質ギャップ分析（P2: 網羅性の事前把握に対応）")

para(doc, "TRMと既存テストを照合し、既存テストで未カバーのテスト要求を種別ごとに集計した（図5）。"
     "なお、既存テストとの照合は既存テストコードの内容からの推定に基づく。")

add_table(doc,
          ["種別", "TRM要求数", "既存テストでカバー済み（推定）", "未カバー（推定）", "未カバー率"],
          [
              ["BR", "55", "26", "29", "52.7%"],
              ["EC", "27", "8", "19", "70.4%"],
              ["BV", "11", "3", "8", "72.7%"],
              ["ER", "3", "1", "2", "66.7%"],
              ["DP", "3", "0", "3", "100%"],
              ["合計", "99", "38", "61", "61.6%"],
          ],
          col_widths=[2.0, 2.5, 4.0, 3.0, 2.5])

# 図5を配置
add_figure(doc, "fig5", "図5: 既存テストの品質ギャップ分析")

para(doc, "種別間の差は顕著であり、DP(100%) > BV(72.7%) > EC(70.4%) > ER(66.7%) > BR(52.7%)の順で未カバー率が高い。"
     "DP（関数間の差異検証・往復変換検証）は既存テストに完全に欠落していた。BV（境界値）とEC（同値クラス）の"
     "未カバー率も70%を超えており、既存の手動テストが分岐の主要パス（BR）に偏っていることが定量的に示された。")

para(doc, "この分析自体が「TRMが存在するからこそ可能」であることに注意すべきである。TRMがなければ、既存テストの"
     "品質ギャップは「テストが少ない」という漠然とした認識にとどまり、「何の観点が欠けているか」を構造的に把握する"
     "ことはできない。これはP2（網羅性の事前把握）への解決策として、TRMがテスト「前」に品質ギャップを種別ごとに"
     "特定できることを示す。")

# 3.7 カバレッジ精密分析結果
subheading(doc, "3.7 カバレッジ精密分析結果")

para(doc, "監査済みTRMから生成された全248件のテストスイートに対し、C0/C1/MC-DC/ループ境界の各基準で"
     "カバレッジを精密分析した結果を表6に示す。")

para(doc, "表6: 関数ごとのカバレッジ達成状況", sz=9, bold=True, after=Pt(2))
add_table(doc,
          ["関数", "コード行", "C0（ステートメント）", "C1（ブランチ）", "MC/DC", "ループ境界"],
          [
              ["GetDateTimeFormat", "24行", "19/19 100%", "24/24 100%", "N/A（原子条件のみ）", "3/3 100%"],
              ["ParseVersion", "33行", "33/33 100%", "31/31 100%", "12/12 100%", "15/15 100%"],
              ["CompareVersion", "4行", "3/3 100%", "3/3 100%", "N/A", "N/A"],
              ["IsMailAddress", "40行", "40/40 100%", "24/24 100%", "17/18 94%", "9/9 100%"],
              ["WhatKindOfTwoChars", "16行", "16/16 100%", "22/22 100%", "18/18 100%", "N/A"],
              ["WhatKindOfTwoChars4KW", "16行", "16/16 100%", "22/22 100%", "18/18 100%", "N/A"],
              ["Convert_ZeneisuToHaneisu", "14行", "14/14 100%", "10/10 100%", "9/9 100%", "3/3 100%"],
              ["Convert_HaneisuToZeneisu", "14行", "14/14 100%", "10/10 100%", "9/9 100%", "3/3 100%"],
              ["合計", "161行", "155/155 100%", "146/146 100%", "83/84 99%", "33/33 100%"],
          ],
          col_widths=[3.5, 1.5, 2.8, 2.5, 2.5, 2.0],
          fs=7.5)

para(doc, "C0（ステートメントカバレッジ）およびC1（ブランチカバレッジ）は全8関数で100%を達成した。"
     "MC/DCは83/84条件ペアで99%を達成し、唯一の未カバーはIsMailAddressのドメインラベル走査ループ内の"
     "大文字A-Z条件（実質的影響が極めて限定的な補完項目）であった。ループ境界は全33ケースで0回/1回/N回の"
     "テストが存在し100%を達成した。")

para(doc, "理論的最小件数（MC/DC達成に約90件、同値クラス+境界値を含めて約210件）に対し、実際の248件は"
     "推奨範囲（210-260件）内であり、テスト要求1件あたり平均2.5件のテストケースは境界値・組み合わせ・"
     "エラーパスを含む包括的なテストスイートとして適切な展開倍率である。")

# 3.8 可読性分析
subheading(doc, "3.8 可読性分析")

para(doc, "TRMの可読性（コード非専門者がテスト要求を理解できるか）を種別別・領域別に分析した結果を以下に示す。")

add_table(doc,
          ["種別", "件数", "L1（コード知識不要）", "L2（ドメイン知識のみ）", "L3（コード知識必要）", "可読率(L1+L2)"],
          [
              ["BR", "55", "20", "14", "21", "61.8%"],
              ["EC", "27", "13", "9", "5", "81.5%"],
              ["BV", "11", "3", "5", "3", "72.7%"],
              ["ER", "3", "3", "0", "0", "100.0%"],
              ["DP", "3", "0", "1", "2", "33.3%"],
              ["全体", "99", "38", "27", "34", "65.7%"],
          ],
          col_widths=[2.0, 1.5, 3.0, 3.0, 3.0, 2.5])

add_table(doc,
          ["領域", "件数", "可読率(L1+L2)"],
          [
              ["format.cpp（日時書式・バージョン解析）", "46", "78.3%"],
              ["convert_util.cpp（全角半角英数変換）", "19", "73.7%"],
              ["CWordParse.cpp（メールアドレス・文字種判定）", "34", "41.2%"],
          ],
          col_widths=[7.0, 2.5, 3.0])

para(doc, "全体の約2/3（65.7%）がコード非専門者にも可読であった。一般に理解しやすいドメイン（日時書式、全角半角変換）"
     "では可読率が高く、エディタ固有のロジック（文字種結合判定のECharKind列挙型やマッピング規則等）を扱う関数では"
     "可読率が低い。なお、L3に分類された34件のうち一部は、description文の改善（例: 「ラテン→CSYMマッピング」を"
     "「ラテン系文字がアルファベットと同種として扱われること」に言い換え）によりL2に引き上げ可能と考えられる。")

para(doc, "本分析は著者による分類であり、実際の非開発者による理解度評価は3.8節の制約として留意する必要がある。")

# ================================================================
# 4. 結論
# ================================================================
heading(doc, "4. 結論", lv=2)

# 4.1 まとめ
subheading(doc, "4.1 まとめ")

para(doc, "本研究では、ソースコードの静的解析からテスト要求モデル（TRM）をYAML形式で生成し、TRM網羅性監査を経て"
     "テスト設計書を作成し、TRMを入力として大規模言語モデルに単体テストを生成させる手法を提案・実証した。"
     "OSSのC++コード8関数（約385行）に適用した結果:")

bold_para(doc, [
    ("1. TRM生成と網羅性監査", True),
    (": 5種別99件のテスト要求を初回生成し、独立監査により64件の漏れを特定して163件に拡充した。"
     "初回の分岐網羅率87%は監査後に推定97%に改善された。この結果はLLMによるテスト設計が主要分岐パスの"
     "抽出に有効である一方、境界値（2.6倍増）とエラーパス（4.3倍増）の抽出に構造的弱点を持つことを示す。"
     "テスト要求の65.7%がコード非専門者にも可読であった。", False)
])

bold_para(doc, [
    ("2. テスト設計書の生成", True),
    (": TRM網羅性監査結果に基づき、関数ごとの正式なテスト設計書（分岐条件一覧・TRM対応表・カバレッジ達成計画）を"
     "生成した。テスト設計書はテストコード生成の入力であると同時に、テスト設計レビューとトレーサビリティの"
     "基盤として機能する。", False)
])

bold_para(doc, [
    ("3. テスト生成とカバレッジ達成", True),
    (": 監査済みTRMから248件のテストを生成し、全要求をカバー。C0/C1カバレッジ100%、MC/DC 99%、"
     "ループ境界100%を達成した。MC/DCは航空宇宙（DO-178C Level A）や自動車（ISO 26262 ASIL D）で"
     "要求される高水準の基準であり、LLMテスト生成とTRMの組み合わせがこの水準を達成可能であることを示した。", False)
])

bold_para(doc, [
    ("4. 実行精度", True),
    (": 初回PASS率91.7%。FAIL 12件は全てLLMの期待値推論誤り（対象関数バグ0件）。"
     "TRM IDにより原因を5分類に即時特定し、修正後100% PASS。", False)
])

bold_para(doc, [
    ("5. 品質ギャップ発見", True),
    (": 既存テストの未カバー要求を推定61件（61.6%）、種別ごとに特定した。特にDP種別は100%未カバーであり、"
     "TRMによる構造化された品質ギャップの可視化が有効に機能した。", False)
])

para(doc, "TRMの導入により、テスト設計根拠の明示（P1）、網羅性の事前定量管理（P2）、テスト失敗時の構造的原因分類（P3）を実現した。")

# 4.2 今後の課題
subheading(doc, "4.2 今後の課題")

para(doc, "本実験は純粋関数8関数に限定されており、以下の課題が残る。")

bold_para(doc, [
    ("(1) 副作用関数への拡張", True),
    (": 本実験の対象は純粋関数に限定されており、ER/DPが計6件（初回）と少ない。"
     "モック・スタブを要する関数でのTRM生成ではER/DP比率の増加が予測され、TRMスキーマの拡張"
     "（モック条件の記述等）が必要である。TRM網羅性監査の結果、監査後でもER 13件・DP 8件にとどまっており、"
     "副作用関数への適用でこれらの種別の有効性をさらに検証する必要がある。", False)
])

bold_para(doc, [
    ("(2) TRM網羅性監査の自動化", True),
    (": 現在の網羅性監査はソースコードの分岐条件とTRM要求の手動照合に基づいている。"
     "gcov/llvm-cov等のカバレッジツール出力とTRM要求の自動照合により、監査プロセスの効率化と"
     "客観性の向上が期待できる。", False)
])

bold_para(doc, [
    ("(3) 可読性の被験者評価", True),
    (": 本研究の可読率65.7%は著者による分類であり、実際の非開発者（QAエンジニア、プロジェクトマネージャー等）"
     "による理解度評価実験が必要である。", False)
])

bold_para(doc, [
    ("(4) 他言語・他プロジェクトでの再現実験", True),
    (": 本実験はC++の1プロジェクト（sakura-editor）に限定されている。TRMの5種別体系は言語非依存に設計されているが、"
     "Java/Python等の他言語やWebアプリケーション等の異なるドメインでの再現実験により一般化可能性を検証する"
     "必要がある。", False)
])

# ================================================================
# 参考文献
# ================================================================
heading(doc, "参考文献", lv=2)

references = [
    '[1] C. Lemieux, J. P. Inala, S. K. Lahiri, and S. Sen, "CodaMOSA: Escaping Coverage Plateaus in Test Generation with Pre-trained Large Language Models," in Proc. ICSE 2023, pp. 919-931, 2023. https://dl.acm.org/doi/10.1109/icse48619.2023.00085',
    '[2] Y. Chen et al., "ChatUniTest: A Framework for LLM-Based Test Generation," in Companion Proc. FSE 2024, ACM, 2024. https://dl.acm.org/doi/abs/10.1145/3663529.3663801',
    '[3] M. Schafer et al., "An Empirical Evaluation of Using Large Language Models for Automated Unit Test Generation," IEEE Transactions on Software Engineering, vol. 50, 2024. https://ieeexplore.ieee.org/document/10329992/',
    '[4] J. A. Ramos, E. M. Bader, J. Bader, and S. Iyer, "CoverUp: Coverage-Guided LLM-Based Test Generation," arXiv:2403.16218, 2024. https://arxiv.org/abs/2403.16218',
    '[5] Qodo (CodiumAI), "Qodo-Cover: An AI-Powered Tool for Automated Test Generation and Code Coverage Enhancement," GitHub, 2024. https://github.com/Codium-ai/cover-agent',
    '[6] "Enhancing LLM-based Test Generation for Hard-to-Cover Branches via Program Analysis (TELPA)," arXiv:2404.04966, ACM TOSEM, 2024. https://dl.acm.org/doi/10.1145/3748505',
    '[7] B. Chu, Y. Feng, K. Liu, Z. Guo, Y. Zhang, H. Shi, Z. Nan, and B. Xu, "Large Language Models for Unit Test Generation: Achievements, Challenges, and the Road Ahead," arXiv:2511.21382, 2025. https://arxiv.org/abs/2511.21382',
    '[8] N. S. Mathews and M. Nagappan, "Design choices made by LLM-based test generators prevent them from finding bugs," arXiv:2412.14137, 2024. https://arxiv.org/abs/2412.14137',
    '[9] "Test smells in LLM-Generated Unit Tests," arXiv:2410.10628, 2024. https://arxiv.org/abs/2410.10628',
    '[10] "Do LLMs generate test oracles that capture the actual or the expected program behaviour?" arXiv:2410.21136, 2024. https://arxiv.org/abs/2410.21136',
    '[11] G. Adzic, Specification by Example: How Successful Teams Deliver the Right Software, Manning, 2011.',
    '[12] "Multi-Step Generation of Test Specifications using Large Language Models," in Proc. ACL 2025 Industry Track, 2025. https://aclanthology.org/2025.acl-industry.11.pdf',
    '[13] C. Wiecher et al., "Model-based analysis and specification of functional requirements and tests for complex automotive systems," Systems Engineering, Wiley, 2024. https://incose.onlinelibrary.wiley.com/doi/full/10.1002/sys.21748',
    '[14] "Static Analysis as a Feedback Loop: Enhancing LLM-Generated Code Beyond Correctness," arXiv:2508.14419, 2025. https://arxiv.org/abs/2508.14419',
    '[15] 「生成AIを活用したテストパターンマトリックスを用いたテスト生成」, SQiP 2024 一般発表 A4-1, 2024. https://www.juse.jp/sqip/symposium/2024/timetable/files/A4-1_ronbun.pdf',
    '[16] 産業技術総合研究所 (AIST), 「生成AI品質マネジメントガイドライン第1版」, Rev. 1.0.0, 2025年5月. https://www.digiarc.aist.go.jp/publication/aiqm/genaiqm-guidelines-v1.html',
    '[17] IPA 独立行政法人情報処理推進機構, 「テキスト生成AIの導入・運用ガイドライン」, 2024年7月. https://www.ipa.go.jp/jinzai/ics/core_human_resource/final_project/2024/generative-ai-guideline.html',
]

for ref in references:
    # URLを分離して表示
    if "https://" in ref:
        idx = ref.index("https://")
        text_part = ref[:idx].rstrip()
        url_part = ref[idx:]
        p = doc.add_paragraph()
        r1 = p.add_run(text_part + " ")
        set_font(r1, size=8.5)
        r2 = p.add_run(url_part)
        set_font(r2, size=7.5, color=RGBColor(0, 0, 180))
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = Pt(14)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.first_line_indent = Cm(-0.5)
    else:
        p = para(doc, ref, sz=8.5, after=Pt(2))
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        p.paragraph_format.line_spacing = Pt(14)


# ========== 保存 ==========
os.makedirs(os.path.dirname(OUT), exist_ok=True)
doc.save(OUT)
print(f"\n=== DOCX saved: {OUT} ===")
print(f"    File size: {os.path.getsize(OUT):,} bytes")

# ========== 埋め込み画像サイズの検証 ==========
print("\n=== Embedded image verification ===")
with zipfile.ZipFile(OUT, 'r') as zf:
    image_files = [f for f in zf.namelist() if f.startswith('word/media/')]
    print(f"  Total embedded images: {len(image_files)}")
    for img_name in sorted(image_files):
        info = zf.getinfo(img_name)
        print(f"  {img_name}: {info.file_size:,} bytes")
        assert info.file_size >= 70000, f"Embedded image too small (<70KB): {img_name} ({info.file_size} bytes)"

assert len(image_files) == 5, f"Expected 5 images, found {len(image_files)}"
print("\n=== All 5 images verified (>=70KB each) ===")
print("=== DONE ===")
