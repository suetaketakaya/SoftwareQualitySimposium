#!/usr/bin/env python3
"""SQiP 2026 提出用 docx (v4: 理系論文版)"""
import os, cairosvg
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE = "/Users/suetaketakaya/SoftwareQualitySymposium"
FIG = os.path.join(BASE, "drafts", "figures")
OUT = os.path.join(BASE, "report", "submission_v4_2026.docx")

def png(name):
    s = os.path.join(FIG, f"{name}.svg")
    p = os.path.join(FIG, f"{name}.png")
    cairosvg.svg2png(url=s, write_to=p, scale=2.0)
    return p

def h(doc, text, lv=2):
    x = doc.add_heading(text, level=lv)
    for r in x.runs:
        r.font.color.rgb = RGBColor(0,0,0)
        r.font.size = Pt(12 if lv==2 else 11)
    x.paragraph_format.space_before = Pt(6)
    x.paragraph_format.space_after = Pt(3)

def p(doc, text, bold=False, sz=9.5, align=None, after=Pt(3)):
    x = doc.add_paragraph()
    r = x.add_run(text)
    r.font.size = Pt(sz)
    r.bold = bold
    r.font.name = "游明朝"
    if align: x.alignment = align
    x.paragraph_format.space_after = after
    x.paragraph_format.space_before = Pt(0)
    x.paragraph_format.line_spacing = Pt(13.5)
    return x

def code(doc, text, sz=8):
    x = doc.add_paragraph()
    r = x.add_run(text)
    r.font.size = Pt(sz)
    r.font.name = "Consolas"
    x.paragraph_format.space_after = Pt(3)
    x.paragraph_format.space_before = Pt(2)
    x.paragraph_format.left_indent = Cm(0.5)

def tbl(doc, headers, rows, cw=None, fs=8):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i,hdr in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = hdr
        for pp in c.paragraphs:
            pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for rr in pp.runs: rr.bold = True; rr.font.size = Pt(fs)
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = str(val)
            for pp in c.paragraphs:
                for rr in pp.runs: rr.font.size = Pt(fs)
    if cw:
        for i,w in enumerate(cw):
            for row in t.rows: row.cells[i].width = Cm(w)

def fig(doc, path, cap, w=Cm(13)):
    x = doc.add_paragraph()
    x.alignment = WD_ALIGN_PARAGRAPH.CENTER
    x.add_run().add_picture(path, width=w)
    x.paragraph_format.space_after = Pt(1)
    c = doc.add_paragraph(cap)
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in c.runs: r.font.size = Pt(8.5); r.italic = True
    c.paragraph_format.space_after = Pt(4)

def main():
    pngs = {n: png(n) for n in ["fig1-overview","fig2-requirements-breakdown",
                                  "fig3-test-comparison","fig4-execution-results"]}
    doc = Document()
    sec = doc.sections[0]
    sec.page_width=Cm(21); sec.page_height=Cm(29.7)
    sec.top_margin=Cm(1.8); sec.bottom_margin=Cm(1.8)
    sec.left_margin=Cm(2.0); sec.right_margin=Cm(2.0)
    st = doc.styles["Normal"]
    st.font.name="游明朝"; st.font.size=Pt(9.5)
    st.paragraph_format.space_after=Pt(2); st.paragraph_format.line_spacing=Pt(13.5)

    # === PAGE 1: 申込書 ===
    p(doc,"ソフトウェア品質シンポジウム2026　「経験論文」「経験発表」",bold=True,sz=11,align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc,"発表申込／アブストラクト　記入用紙",bold=True,sz=10.5,align=WD_ALIGN_PARAGRAPH.CENTER,after=Pt(10))
    p(doc,"〔発表申込書〕",bold=True,sz=10.5,after=Pt(6))
    p(doc,"タイトル",bold=True,sz=9.5)
    p(doc,"静的解析に基づくテスト要求モデルの自動生成と、それを用いたLLM単体テスト生成の実証評価")
    p(doc,"申込区分",bold=True,sz=9.5); p(doc,"経験発表")
    p(doc,"カテゴリ",bold=True,sz=9.5); p(doc,"・品質管理・テスト技術の観点")
    p(doc,"キーワード",bold=True,sz=9.5)
    p(doc,"ホワイトボックステスト / テスト要求モデル / LLM / 静的解析 / 単体テスト生成 / トレーサビリティ",after=Pt(10))

    # === PAGE 2+: アブストラクト ===
    doc.add_page_break()
    p(doc,"〔アブストラクト記入用紙〕",bold=True,sz=10.5,align=WD_ALIGN_PARAGRAPH.CENTER,after=Pt(6))

    # --- 1. ねらい ---
    h(doc,"1. ねらい")
    p(doc,
      "LLMを活用したソフトウェア開発では、要求仕様からコードを直接生成し、"
      "基本機能の動作確認で品質を評価する運用が広がっている。"
      "しかし実リリースを想定すると、単体テスト・コンポーネントテストの層で"
      "品質を担保する必要がある。"
      "現状のLLMテスト生成手法 [1]-[4] は"
      "カバレッジ向上やコンパイル成功率の改善に注力しているが、"
      "共通して以下の3つの問題を残している。")
    p(doc,
      "(P1) テスト設計根拠の不在: 生成テストが「何の分岐を」「どの同値クラスを」"
      "対象としているか不明であり、テスト設計の妥当性を第三者が検証できない。"
      "(P2) 網羅性の事前把握が不可能: テスト群がソースコードのどの範囲をカバーしているかは"
      "実行後にしか分からず、テスト生成「前」に網羅性を管理する手段がない。"
      "(P3) テスト失敗時の原因特定コスト: 失敗時にテストコードとソースコードを"
      "逐一突合する必要があり、原因分類に時間がかかる。")
    p(doc,
      "本研究は、ソースコードの静的解析結果からホワイトボックステスト要求モデル"
      "（以下TRM）をYAML形式で自動生成し、TRMを入力としてLLMにテストコードを"
      "生成させる手法を提案する。"
      "TRMは分岐条件・同値クラス・境界値をID付きで構造化した中間成果物であり、"
      "テスト生成の「前」にテスト設計の網羅性を定量評価できる。"
      "OSSプロジェクトのC++コード8関数に適用し、TRMの生成精度、テスト生成品質、"
      "テスト失敗時の診断効率を定量評価した。")

    # --- 2. 実施概要 ---
    h(doc,"2. 実施概要")
    p(doc,"提案手法は3段階のパイプラインである（図1）。")
    p(doc,
      "Step 1（対象選定）: 対象リポジトリのソースコード構造を解析し、"
      "(a) 外部依存なし、(b) 分岐構造が明確、(c) 既存テスト存在の3基準で"
      "テスト対象関数を選定する。")
    p(doc,
      "Step 2（TRM生成）: 各関数のソースコードから分岐条件を抽出し、"
      "5種別（BR: 分岐網羅、EC: 同値クラス、BV: 境界値、ER: エラーパス、"
      "DP: 依存切替）でテスト要求をYAML形式で定義する。"
      "各要求にはID・説明・優先度・ソースコード行参照を付与する。"
      "以下はParseVersion関数に対するTRMの抜粋である。")

    code(doc,
      'id: "BR-20"\n'
      'type: branch_coverage\n'
      'description: "alpha修飾子の完全一致分岐を通過する"\n'
      'priority: high\n'
      'source_ref: "BC-02-10: strncmp(p,\\"alpha\\",5)==0"')

    p(doc,
      "Step 3（テスト生成）: TRMのYAMLと対象ソースコードをLLMに入力し、"
      "Google Test準拠のテストコードを生成させる。"
      "各テストケースにTRMのIDをコメントで記載する。")

    code(doc,
      'TEST(ParseVersion, AlphaModifier) {\n'
      '    // BR-20: alpha修飾子の完全一致分岐\n'
      '    UINT32 val = ParseVersion(L"2.4.1alpha");\n'
      '    EXPECT_EQ(val, 0x82848120);\n'
      '}')

    fig(doc, pngs["fig1-overview"], "図1: 提案手法の全体構成（3段階パイプライン）", Cm(14))

    p(doc,
      "実験条件: sakura-editor/sakura（C++, OSS）の3領域8関数（約385行）に適用した。"
      "実行環境はmacOS + Apple clang 16.0 + Google Test 1.17 + Windows互換ヘッダ。"
      "比較対象は既存の手動テスト（約45件）とした。")

    # --- 3. 実施結果 ---
    h(doc,"3. 実施結果")

    p(doc,
      "3.1 TRM生成結果: 8関数に対し計99件のテスト要求を定義した（図2）。"
      "種別の内訳はBR 55件、EC 27件、BV 11件、ER 3件、DP 3件。"
      "約385行のコードに対し平均3.9行/要求の密度でテスト要求を抽出した。",
      sz=9.5)
    fig(doc, pngs["fig2-requirements-breakdown"], "図2: テスト要求の種別内訳（N=99）", Cm(9))

    p(doc,
      "3.2 テスト生成・実行結果: TRMを入力として145件のテストコードを生成した（図3）。"
      "テスト要求カバー率は100%。既存テスト45件の約3.2倍。"
      "コンパイルは3/3ファイル成功。初回テストPASS率91.7%（133/145件）。"
      "FAIL 12件は全てLLMの期待値推論誤りであり、対象関数の実装バグは0件（図4）。"
      "TRM IDから原因を5分類に即時特定し、修正後に全145件PASS。",
      sz=9.5)
    fig(doc, pngs["fig3-test-comparison"], "図3: 既存テストと生成テストの比較", Cm(11))
    fig(doc, pngs["fig4-execution-results"], "図4: コンパイル・実行結果", Cm(11))

    p(doc,"3.3 FAIL原因の分類（TRM IDによる構造的診断）:",bold=True,sz=9)
    tbl(doc,
      ["原因分類","件数","対応TRM ID","技術的詳細"],
      [
        ["コンポーネント分割の誤解","5","BR-20〜23",'"2.4.1alpha"を1コンポーネントと予測。実際は"1"と"alpha"が別処理'],
        ["数字グルーピングの誤解","1","EC-08",'"1234"を1桁ずつと予測。実際は2桁ずつ区切り'],
        ["文字数カウントミス","2","EC-13,14","wcslen=18をnull込み19と誤算"],
        ["API呼び出しバグ","1","EC-16","offsetとバッファポインタの二重シフト"],
        ["マッピング戻り値の誤解","3","BR-39〜41","変数マッピング後の返却値を元値と誤認"],
      ],
      cw=[3.5,1.0,2.5,7.0], fs=7.5)

    p(doc,
      "3.4 既存テストとの品質ギャップ: TRMと既存テストを照合した結果、"
      "全要求の61.6%（61/99件）が既存テストで未カバーであった。"
      "種別ごとの未カバー率はDP 100%、BV 73%、EC 70%、ER 67%、BR 53%。"
      "特にDP（関数間の差異検証・往復変換検証）は既存テストに完全に欠落していた。"
      "TRMにより、テスト生成「前」に品質ギャップを種別ごとに特定できることが示された。",
      sz=9.5)

    # --- 4. 結論 ---
    h(doc,"4. 結論")
    p(doc,
      "本研究では、ソースコードの静的解析からTRMをYAML形式で生成し、"
      "TRMを入力としてLLMに単体テストを生成させる手法を提案・実証した。"
      "C++コード8関数（約385行）への適用結果: "
      "(1) 5種別99件のTRMを構造化定義（平均3.9行/要求）。"
      "(2) TRMからの145件テストが全要求カバー。既存テスト45件の3.2倍を体系的に生成。"
      "(3) 初回PASS率91.7%。FAIL 12件はTRM IDで5分類に即時特定、修正後100% PASS。"
      "(4) 既存テストの未カバー要求61件（61.6%）を種別ごとに特定（DP種別は100%未カバー）。"
      "TRMの導入により、テスト設計根拠の明示(P1)、網羅性の事前管理(P2)、"
      "テスト失敗の構造的原因分類(P3)を実現した。"
      "今後の課題は、副作用関数へのスタブ化拡張、C0/C1カバレッジとの相関分析、"
      "他言語での再現実験である。")

    # 参考文献
    p(doc,"",sz=3,after=Pt(1))
    p(doc,"参考文献",bold=True,sz=8.5,after=Pt(1))
    p(doc,
      '[1] C. Lemieux et al., "CodaMOSA," ICSE 2023.  '
      '[2] Y. Chen et al., "ChatUniTest," FSE 2024.  '
      '[3] M. Schafer et al., "LLMs for Unit Test Generation," IEEE TSE 2024.  '
      '[4] J. Ramos et al., "CoverUp," arXiv:2403.16218, 2024.  '
      '[5] B. Chu et al., "LLMs for Unit Test Generation: Survey," arXiv:2511.21382, 2025.',
      sz=7.5,after=Pt(0))

    doc.save(OUT)
    print(f"Saved: {OUT}")

if __name__ == "__main__":
    main()
