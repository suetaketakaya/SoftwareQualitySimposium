#!/usr/bin/env python3
"""generate_submission_v6.py — paper-draft-full.md から docx 投稿原稿を生成

次稿論文（TRM v3.1 + EN種別 + 可視化レイヤ）の docx 投稿原稿を、
Markdown ドラフトから機械的に生成する。

v6 の位置づけ（v4/v5 との差分）:
- v4/v5 は SQiP 2026 提出用にハードコードされた個別スクリプト
- v6 は knowledge/paper-draft-full.md を唯一の情報源として動的生成
- Markdown 更新のみで docx が再生成され、内容と原稿の整合を保つ

Usage:
  python scripts/generate_submission_v6.py \\
    --input knowledge/paper-draft-full.md \\
    --output report/submission_v6_2026.docx
"""
from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor


# ---------------------------------------------------------------------------
# Markdown パース
# ---------------------------------------------------------------------------

def parse_markdown(md_text: str) -> list[dict]:
    """Markdown を docx 要素リストに変換する"""
    elements: list[dict] = []
    lines = md_text.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]

        # 引用・水平線・YAML フロントマターはスキップ
        if line.strip().startswith(">") or line.strip() in ("---", "___", "***"):
            i += 1
            continue

        # 見出し
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            elements.append({"type": f"h{level}", "text": text})
            i += 1
            continue

        # コードブロック
        if line.startswith("```"):
            lang = line[3:].strip()
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # closing ```

            if lang == "mermaid":
                elements.append({
                    "type": "figure_placeholder",
                    "text": "[図: Mermaid 図版 — paper-figures-plan.md 参照]",
                })
            else:
                elements.append({
                    "type": "code",
                    "text": "\n".join(code_lines),
                    "lang": lang,
                })
            continue

        # 表（header + separator + rows）
        if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s\-|:]+\|?$", lines[i+1]):
            rows = []
            while i < len(lines) and lines[i].startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                rows.append(cells)
                i += 1
            if len(rows) >= 2 and all("-" in c or ":" in c for c in rows[1]):
                rows = [rows[0]] + rows[2:]
            elements.append({"type": "table", "rows": rows})
            continue

        # ブレット
        if re.match(r"^\s*[-*]\s+", line):
            items = []
            while i < len(lines) and re.match(r"^\s*[-*]\s+", lines[i]):
                items.append(re.sub(r"^\s*[-*]\s+", "", lines[i]))
                i += 1
            elements.append({"type": "bullet_list", "items": items})
            continue

        # 番号付きリスト
        if re.match(r"^\s*\d+\.\s+", line):
            items = []
            while i < len(lines) and re.match(r"^\s*\d+\.\s+", lines[i]):
                items.append(re.sub(r"^\s*\d+\.\s+", "", lines[i]))
                i += 1
            elements.append({"type": "numbered_list", "items": items})
            continue

        # 通常の段落
        if line.strip():
            elements.append({"type": "paragraph", "text": line.strip()})
        i += 1

    return elements


# ---------------------------------------------------------------------------
# インライン処理
# ---------------------------------------------------------------------------

def clean_inline(text: str) -> str:
    """docx table 等で使う plain text 変換"""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)
    return text


def parse_inline_tokens(text: str) -> list[tuple[str, dict]]:
    """インライン記法を (text, format) のトークン列に"""
    tokens: list[tuple[str, dict]] = []
    pattern = re.compile(r"(\*\*.+?\*\*|`.+?`|\*.+?\*|\[.+?\]\(.+?\))")
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            tokens.append((text[last:m.start()], {}))
        matched = m.group(0)
        if matched.startswith("**"):
            tokens.append((matched[2:-2], {"bold": True}))
        elif matched.startswith("`"):
            tokens.append((matched[1:-1], {"code": True}))
        elif matched.startswith("*"):
            tokens.append((matched[1:-1], {"italic": True}))
        elif matched.startswith("["):
            link_text = re.match(r"\[(.+?)\]", matched).group(1)
            tokens.append((link_text, {}))
        last = m.end()
    if last < len(text):
        tokens.append((text[last:], {}))
    return tokens


# ---------------------------------------------------------------------------
# docx 出力
# ---------------------------------------------------------------------------

def add_title_page(doc: Document):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("テスト要求モデルによるホワイトボックス品質情報の非エンジニア共有")
    r.font.size = Pt(16)
    r.bold = True
    r.font.name = "游ゴシック"

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("4種の図式自動生成を用いた AI駆動開発時代の品質可視化")
    r.font.size = Pt(13)
    r.font.name = "游ゴシック"

    doc.add_paragraph()
    anon = doc.add_paragraph()
    anon.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = anon.add_run("[著者情報 — 匿名査読のため空欄]")
    r.italic = True
    r.font.size = Pt(10)

    doc.add_paragraph()
    ver = doc.add_paragraph()
    ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = ver.add_run("Draft v6 — 2026-04-21")
    r.italic = True
    r.font.size = Pt(9)

    doc.add_page_break()


def add_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    for tok_text, fmt in parse_inline_tokens(text):
        run = p.add_run(tok_text)
        if fmt.get("bold"):
            run.bold = True
        if fmt.get("italic"):
            run.italic = True
        if fmt.get("code"):
            run.font.name = "Consolas"
            run.font.size = Pt(9)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.3


def add_code_block(doc: Document, code: str):
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(3)


def add_figure_placeholder(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def add_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Light Grid Accent 1"
    for ri, row in enumerate(rows):
        for ci, cell in enumerate(row):
            if ci < len(t.rows[ri].cells):
                t.rows[ri].cells[ci].text = clean_inline(cell)
                for para in t.rows[ri].cells[ci].paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)
                        if ri == 0:
                            run.bold = True


def render_elements(doc: Document, elements: list[dict]):
    """パース済み要素を docx に書き込む。執筆メタ情報セクションはスキップ"""
    skip = False
    for el in elements:
        t = el["type"]
        if t == "h2" and "執筆状況" in el.get("text", ""):
            skip = True
        if skip:
            continue

        text = el.get("text", "")

        # トップレベルのタイトル・副題・ドラフトメタ情報はタイトルページで表示済
        if t == "h1":
            if "テスト要求モデルによる" in text or "章立て" in text:
                continue
            doc.add_heading(clean_inline(text), level=0)
        elif t == "h2":
            doc.add_heading(clean_inline(text), level=1)
        elif t == "h3":
            doc.add_heading(clean_inline(text), level=2)
        elif t == "h4":
            doc.add_heading(clean_inline(text), level=3)
        elif t in ("h5", "h6"):
            doc.add_heading(clean_inline(text), level=4)
        elif t == "paragraph":
            add_paragraph(doc, text)
        elif t == "bullet_list":
            for item in el["items"]:
                doc.add_paragraph(clean_inline(item), style="List Bullet")
        elif t == "numbered_list":
            for item in el["items"]:
                doc.add_paragraph(clean_inline(item), style="List Number")
        elif t == "table":
            add_table(doc, el["rows"])
        elif t == "code":
            add_code_block(doc, el["text"])
        elif t == "figure_placeholder":
            add_figure_placeholder(doc, el["text"])


def apply_document_defaults(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "游明朝"
    style.font.size = Pt(10)

    for level in ("Heading 1", "Heading 2", "Heading 3", "Heading 4"):
        try:
            doc.styles[level].font.name = "游ゴシック"
        except KeyError:
            pass

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description=__doc__.splitlines()[0])
    parser.add_argument("--input", required=True, type=Path,
                        help="Markdown paper draft (paper-draft-full.md)")
    parser.add_argument("--output", required=True, type=Path,
                        help="Output .docx file path")
    args = parser.parse_args()

    md_text = args.input.read_text(encoding="utf-8")
    elements = parse_markdown(md_text)

    doc = Document()
    apply_document_defaults(doc)
    add_title_page(doc)
    render_elements(doc, elements)

    args.output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(args.output))

    h = sum(1 for e in elements if e["type"].startswith("h"))
    p = sum(1 for e in elements if e["type"] == "paragraph")
    tbl = sum(1 for e in elements if e["type"] == "table")
    code = sum(1 for e in elements if e["type"] == "code")
    fig = sum(1 for e in elements if e["type"] == "figure_placeholder")

    print(f"Generated: {args.output}")
    print(f"  Headings={h} Paragraphs={p} Tables={tbl} Code={code} Figures={fig}")


if __name__ == "__main__":
    main()
