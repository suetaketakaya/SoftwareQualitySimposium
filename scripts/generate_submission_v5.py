#!/usr/bin/env python3
"""SQiP 2026 提出用 docx (v5: 理系論文・詳細版・matplotlib PNG直接使用)"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE = "/Users/suetaketakaya/SoftwareQualitySymposium"
FIG = os.path.join(BASE, "drafts", "figures")
OUT = os.path.join(BASE, "report", "submission_v5_2026.docx")

# matplotlib生成済みPNGのパスを直接参照（cairosvg不使用）
PNGS = {
    "fig1": os.path.join(FIG, "fig1-overview.png"),
    "fig2": os.path.join(FIG, "fig2-requirements-breakdown.png"),
    "fig3": os.path.join(FIG, "fig3-test-comparison.png"),
    "fig4": os.path.join(FIG, "fig4-execution-results.png"),
}

def heading(doc, text, lv=2):
    x = doc.add_heading(text, level=lv)
    for r in x.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.size = Pt(13 if lv == 2 else 11)
    x.paragraph_format.space_before = Pt(10)
    x.paragraph_format.space_after = Pt(4)

def subheading(doc, text):
    x = doc.add_paragraph()
    r = x.add_run(text)
    r.font.size = Pt(10.5)
    r.bold = True
    r.font.name = "游明朝"
    x.paragraph_format.space_before = Pt(8)
    x.paragraph_format.space_after = Pt(2)

def para(doc, text, sz=10, after=Pt(4), bold=False, align=None):
    x = doc.add_paragraph()
    r = x.add_run(text)
    r.font.size = Pt(sz)
    r.bold = bold
    r.font.name = "游明朝"
    if align:
        x.alignment = align
    x.paragraph_format.space_after = after
    x.paragraph_format.space_before = Pt(0)
    x.paragraph_format.line_spacing = Pt(16)
    return x

def code_block(doc, text):
    x = doc.add_paragraph()
    r = x.add_run(text)
    r.font.size = Pt(8.5)
    r.font.name = "Consolas"
    x.paragraph_format.space_after = Pt(4)
    x.paragraph_format.space_before = Pt(2)
    x.paragraph_format.left_indent = Cm(0.8)
    x.paragraph_format.line_spacing = Pt(13)

def table(doc, headers, rows, col_widths=None, fs=8.5):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hdr in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = hdr
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(fs)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            c.text = str(val)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(fs)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)

def figure(doc, png_path, caption, width=Cm(14)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(png_path, width=width)
    p.paragraph_format.space_after = Pt(2)
    c = doc.add_paragraph(caption)
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in c.runs:
        r.font.size = Pt(9)
        r.italic = True
    c.paragraph_format.space_after = Pt(6)


def main():
    # PNGファイルの存在確認
    for k, path in PNGS.items():
        assert os.path.exists(path), f"PNG not found: {path}"
        size = os.path.getsize(path)
        assert size > 50000, f"PNG too small (likely broken): {path} ({size} bytes)"

    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    st = doc.styles["Normal"]
    st.font.name = "游明朝"
    st.font.size = Pt(10)
    st.paragraph_format.space_after = Pt(4)
    st.paragraph_format.line_spacing = Pt(16)

    # ======================================================================
    # PAGE 1: 発表申込書
    # ======================================================================
    para(doc, "ソフトウェア品質シンポジウム2026　「経験論文」「経験発表」",
         bold=True, sz=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "発表申込／アブストラクト　記入用紙",
         bold=True, sz=11, align=WD_ALIGN_PARAGRAPH.CENTER, after=Pt(14))

    para(doc, "〔発表申込書〕", bold=True, sz=11, after=Pt(8))

    para(doc, "タイトル", bold=True, sz=10)
    para(doc, "静的解析に基づくテスト要求モデルの自動生成と、"
         "それを用いたLLM単体テスト生成の実証評価", sz=10.5)

    para(doc, "申込区分", bold=True, sz=10)
    para(doc, "経験発表", sz=10.5)

    para(doc, "カテゴリ", bold=True, sz=10)
    para(doc, "・品質管理・テスト技術の観点", sz=10.5)

    para(doc, "キーワード", bold=True, sz=10)
    para(doc, "ホワイトボックステスト / テスト要求モデル / 大規模言語モデル（LLM） / "
         "静的解析 / 単体テスト生成 / トレーサビリティ", sz=10.5, after=Pt(14))

    # ======================================================================
    # PAGE 2+: アブストラクト
    # ======================================================================
    doc.add_page_break()
    para(doc, "〔アブストラクト記入用紙〕", bold=True, sz=11,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=Pt(10))

    # ------------------------------------------------------------------
    # 1. ねらい
    # ------------------------------------------------------------------
    heading(doc, "1. ねらい")

    subheading(doc, "1.1 背景: LLM活用開発における単体テストの課題")

    para(doc,
        "大規模言語モデル（LLM）を活用したソフトウェア開発が急速に広がっている。"
        "現在の典型的な運用は、自然言語で記述した要求仕様をLLMに与え、"
        "実装コードを直接生成させるものである。"
        "この開発スタイルでは、生成されたコードの品質評価が"
        "「要求どおりに基本機能が動作するか」の確認に偏りがちであり、"
        "単体テスト・コンポーネントテストの層が構造的に手薄になる。")

    para(doc,
        "しかし、生成されたコードを実プロダクトとしてリリースするためには、"
        "従来のコードベース開発と同等以上の品質保証が必要である。"
        "特に、各関数が仕様どおりに動作すること、"
        "分岐条件や境界値が正しく処理されること、"
        "異常系でも安全に動作することを、"
        "単体テストによって体系的に検証する必要がある。")

    subheading(doc, "1.2 既存のLLMテスト生成手法の限界")

    para(doc,
        "LLMによるテスト自動生成は近年活発に研究されている。"
        "CodaMosa [1] はサーチベーステストとLLMを組み合わせてカバレッジ向上を図り、"
        "ChatUniTest [2] はLLMによるテスト生成とコンパイルエラーの反復修正を行い、"
        "TestPilot [3] はJavaScriptの単体テスト自動生成を実現し、"
        "CoverUp [4] はカバレッジフィードバックによるテスト改善を提案している。"
        "これらの手法はカバレッジ向上やコンパイル成功率の改善に成果を上げているが、"
        "共通して以下の3つの問題を残している。")

    para(doc,
        "(P1) テスト設計根拠の不在: LLMが生成したテストが「何の分岐を」"
        "「どの同値クラスを」対象としているか不明であり、"
        "テスト設計の妥当性を第三者が検証できない。"
        "Mathews & Nagappan [5] は、LLMテスト生成ツールの設計上の選択が"
        "バグ発見を妨げていることを指摘しており、"
        "LLM生成テストにはテストスメルの頻出 [7] や"
        "テストオラクルの正確性50%以下 [8] という品質課題も報告されている。")

    para(doc,
        "(P2) 網羅性の事前把握が不可能: テスト群がソースコードのどの範囲を"
        "カバーしているかは、テストをコンパイル・実行してカバレッジツールを通すまで分からない。"
        "テスト生成の「前」に、何をテストすべきかを設計・管理する手段がない。")

    para(doc,
        "(P3) テスト失敗時の原因特定コスト: テストが失敗した場合、"
        "生成テストのアサーション値とソースコードの実装を逐一突合して"
        "原因を特定する必要がある。テスト数が増えるほどこのコストは増大する。")

    subheading(doc, "1.3 本研究の目的と提案手法")

    para(doc,
        "本研究は、上記の3つの問題に対し、"
        "ソースコードの静的解析結果からホワイトボックステスト要求モデル"
        "（Test Requirement Model、以下TRM）をYAML形式で自動生成し、"
        "TRMを入力としてLLMにテストコードを生成させる手法を提案する。")

    para(doc,
        "TRMは、対象関数の分岐条件・同値クラス・境界値を"
        "一意のID付きで構造化した中間成果物である。"
        "TRMを導入することで、"
        "(P1に対し) 各テストケースがどのテスト要求に対応するかが明示され、テスト設計根拠が追跡可能になる。"
        "(P2に対し) テスト生成の「前」に、テスト設計の網羅性を定量的に評価できる。"
        "(P3に対し) テスト失敗時に、TRM IDから原因の所在を即座に特定できる。")

    para(doc,
        "本研究では、OSSプロジェクトのC++コード8関数（約385行）に"
        "提案手法を適用し、TRMの生成結果、テスト生成品質、"
        "テスト失敗時の診断効率を定量的に評価した。")

    # ------------------------------------------------------------------
    # 2. 実施概要
    # ------------------------------------------------------------------
    heading(doc, "2. 実施概要")

    subheading(doc, "2.1 提案手法の3段階パイプライン")

    para(doc,
        "提案手法は、図1に示す3段階のパイプラインで構成される。")

    para(doc,
        "Step 1（対象選定）: 対象リポジトリのソースコード構造をLLMが解析し、"
        "テスト対象関数を選定する。選定基準は "
        "(a) 外部依存がない純粋関数であること、"
        "(b) if/else/switchによる明確な分岐構造を持つこと、"
        "(c) 既存のテストが存在し比較評価が可能であること、の3点とした。"
        "人間が選定結果をレビューし、承認する。")

    para(doc,
        "Step 2（TRM生成）: 選定した各関数のソースコードから分岐条件を抽出し、"
        "テスト要求をYAML形式で定義する。テスト要求は以下の5種別で体系的に分類する。")

    table(doc,
        ["種別", "記号", "定義", "導出方法"],
        [
            ["分岐網羅", "BR", "if/else/switchの各分岐パス", "制御フロー解析"],
            ["同値クラス", "EC", "入力パラメータの同値分割", "関数シグネチャ+分岐条件"],
            ["境界値", "BV", "同値クラス境界の上限・下限", "EC要求の境界条件"],
            ["エラーパス", "ER", "異常系・エラー処理パス", "ガード条件・例外処理"],
            ["依存切替", "DP", "関数間の依存関係パス", "呼び出し関係・類似関数の差異"],
        ],
        col_widths=[2.0, 1.0, 4.0, 4.0], fs=8.5)

    para(doc,
        "各テスト要求には、一意のID（例: BR-20）、自然言語による説明、"
        "優先度（high/medium/low）、根拠となるソースコードの行番号参照を付与する。"
        "以下に、バージョン文字列解析関数ParseVersionに対するTRMの抜粋を示す。")

    code_block(doc,
        '- id: "BR-20"\n'
        '  type: branch_coverage\n'
        '  description: "alpha修飾子の完全一致分岐を通過する"\n'
        '  priority: high\n'
        '  source_ref: "BC-02-10: strncmp(p, \\"alpha\\", 5) == 0"\n'
        '  expected_behavior: "nShift = -0x60 を設定し、ポインタを5文字進める"')

    para(doc,
        "このYAMLエントリは、「ParseVersion関数の中で、入力文字列中に"
        "'alpha'という修飾子が現れた場合の分岐が正しく動作するか」"
        "というテスト要求を表している。人間はこのYAMLをレビューし、"
        "テスト設計の過不足を修正する。")

    para(doc,
        "Step 3（テストコード生成）: 承認されたTRMのYAMLと対象関数のソースコードを"
        "LLMに入力し、Google Test準拠のテストコードを生成させる。"
        "各テストケースにはTRMのIDをコメントで明記し、"
        "テスト要求とテストケースの間の追跡性（トレーサビリティ）を確保する。"
        "以下に、上記BR-20に対応する生成テストコードを示す。")

    code_block(doc,
        'TEST(ParseVersion, AlphaModifier) {\n'
        '    // BR-20: alpha修飾子の完全一致分岐を通過する\n'
        '    UINT32 val = ParseVersion(L"2.4.1alpha");\n'
        '    EXPECT_EQ(val, static_cast<UINT32>(0x82848120));\n'
        '}')

    figure(doc, PNGS["fig1"],
           "図1: 提案手法の全体構成（3段階パイプライン）", Cm(14.5))

    subheading(doc, "2.2 実験条件")

    para(doc,
        "実験対象として、オープンソースのC++プロジェクト "
        "sakura-editor/sakura（Windows向け日本語テキストエディタ、"
        "Google Testによる既存テスト約45件が整備済み）を使用した。"
        "Step 1により、以下の3領域8関数（合計約385行）を選定した。")

    table(doc,
        ["領域", "対象関数", "コード行数", "主な処理内容"],
        [
            ["日時書式・\nバージョン解析", "GetDateTimeFormat\nParseVersion\nCompareVersion", "約115行", "日時書式変換、バージョン\n文字列の解析・比較"],
            ["メールアドレス・\n文字種判定", "IsMailAddress\nWhatKindOfTwoChars\nWhatKindOfTwoChars4KW", "約180行", "メールアドレス検証、\n日本語文字種の結合判定"],
            ["全角半角変換", "Convert_ZeneisuToHaneisu\nConvert_HaneisuToZeneisu", "約90行", "全角英数字⇔半角英数字\nの相互変換"],
        ],
        col_widths=[2.5, 4.0, 1.8, 4.0], fs=8)

    para(doc,
        "選定過程では、GUI依存（ウィンドウ操作を含む関数）、"
        "ファイルI/O依存、グローバル状態依存の11関数を候補から除外した。"
        "除外理由は、スタブ化が必要となり実験の複雑さが増すためである。"
        "実行環境はmacOS（Apple clang 16.0）+ Google Test 1.17とし、"
        "Windows固有の型定義（SYSTEMTIME、BOOL等）は互換ヘッダで代替した。"
        "比較対象として、既存の手動テスト（約45件）を使用した。")

    # ------------------------------------------------------------------
    # 3. 実施結果
    # ------------------------------------------------------------------
    heading(doc, "3. 実施結果")

    subheading(doc, "3.1 TRMの生成結果")

    para(doc,
        "Step 2により、8関数に対して合計99件のテスト要求を定義した。"
        "図2に種別ごとの内訳を示す。"
        "分岐網羅（BR）が55件（55.6%）と最も多く、"
        "次いで同値クラス（EC）27件（27.3%）、境界値（BV）11件（11.1%）、"
        "エラーパス（ER）3件（3.0%）、依存切替（DP）3件（3.0%）であった。")

    para(doc,
        "BR要求が過半数を占めるのは、ソースコードのif/else/switch分岐から"
        "直接導出されるためである。"
        "ER・DPが少ないのは、対象がいずれも純粋関数であり、"
        "エラーハンドリングや関数間の複雑な依存関係が限定的であることによる。"
        "約385行のソースコードに対し、平均3.9行あたり1件のテスト要求を抽出した "
        "（関数の分岐密度が高いほど要求密度も高くなる傾向がある）。")

    figure(doc, PNGS["fig2"],
           "図2: テスト要求の種別内訳（N=99）", Cm(10))

    subheading(doc, "3.2 テスト生成結果と既存テストとの比較")

    para(doc,
        "Step 3により、TRMの99件の要求を入力として、"
        "合計145件のGoogle Testテストコードを生成した。"
        "テスト要求カバー率は100%（全99件に対応するテストケースが存在）であった。"
        "テストケース数（145件）がテスト要求数（99件）を上回るのは、"
        "1つの境界値要求を上限・下限の別テストケースで検証する設計や、"
        "修飾子のタイプミス入力に対する補強テスト等を含むためである。")

    para(doc,
        "図3に、既存の手動テスト（約45件）と提案手法による生成テスト（145件）の"
        "領域別の比較を示す。全領域で既存テストを上回り、"
        "全体として約3.2倍のテストケースを体系的に生成した。")

    figure(doc, PNGS["fig3"],
           "図3: 既存テストと生成テストの領域別比較", Cm(12))

    subheading(doc, "3.3 コンパイル・実行結果")

    para(doc,
        "生成した145件のテストコードをmacOS環境でコンパイル・実行した。"
        "結果を図4および以下の表に示す。")

    table(doc,
        ["指標", "値"],
        [
            ["コンパイル成功", "3/3ファイル（100%）"],
            ["初回テストPASS", "133/145件（91.7%）"],
            ["初回テストFAIL", "12件（8.3%）"],
            ["FAIL原因: 対象関数の実装バグ", "0件"],
            ["FAIL原因: LLMの期待値推論誤り", "12件（100%）"],
            ["期待値修正後のPASS", "145/145件（100%）"],
        ],
        col_widths=[6.0, 5.0], fs=9)

    para(doc,
        "12件のFAILは全て、LLMがテストの期待値（Expected値）を"
        "誤って推論したものであった。"
        "対象関数自体の実装バグは1件も検出されなかった。"
        "期待値を修正した後に再実行したところ、全145件がPASSした。")

    figure(doc, PNGS["fig4"],
           "図4: コンパイル・実行結果（初回→修正後）", Cm(12))

    subheading(doc, "3.4 FAIL 12件の原因分析: TRMによる構造的診断")

    para(doc,
        "本実験の重要な知見として、12件のFAILに対する原因特定における"
        "TRMの効果を報告する。"
        "各テストケースにはTRM ID（例: BR-20）がコメントで付与されているため、"
        "失敗したテストがどのテスト要求に対応するかを即座に特定でき、"
        "TRMの説明文と根拠情報から原因の所在を構造的に分類できた。"
        "以下にFAIL 12件の原因分類を示す。")

    table(doc,
        ["原因分類", "件数", "対応TRM ID", "技術的詳細"],
        [
            ["コンポーネント分割\nの誤解", "5", "BR-20〜23",
             '"2.4.1alpha"をLLMは1コンポーネントと予測。\n'
             '実際は"1"と"alpha"が別のイテレーションで処理される。'],
            ["数字グルーピング\nの誤解", "1", "EC-08",
             '"1234"を1桁ずつと予測。実際は2桁ずつ区切り。'],
            ["文字数カウント\nミス", "2", "EC-13, 14",
             'wcslen("test@example.co.jp")=18を19と誤算\n'
             '（null終端を含めてカウント）。'],
            ["API呼び出しバグ", "1", "EC-16",
             'IsMailAddressの引数で、offsetとバッファ\n'
             'ポインタを二重にシフトした。'],
            ["マッピング戻り値\nの誤解", "3", "BR-39〜41",
             'WhatKindOfTwoCharsで、変数のマッピング後の\n'
             '返却値を、マッピング前の元値と誤認。'],
        ],
        col_widths=[2.5, 0.8, 2.0, 7.0], fs=8)

    para(doc,
        "TRMがない場合、これらの原因特定には、145件のテストコードの"
        "アサーション値とC++ソースコードの分岐条件を逐一突合する必要がある。"
        "TRMがある場合は、失敗テストのコメントに記載されたTRM ID（例: BR-20）から、"
        "対応するテスト要求の説明（「alpha修飾子の完全一致分岐」）と"
        "ソースコード参照（「strncmp(p, \"alpha\", 5) == 0」）を即座に参照でき、"
        "原因を体系的に分類できた。"
        "これはP3（テスト失敗時の原因特定コスト）への直接的な解決策である。")

    subheading(doc, "3.5 既存テストの品質ギャップ分析")

    para(doc,
        "TRMの99件を基準として既存テスト（約45件）との照合を行い、"
        "既存テストでカバーされていないテスト要求を種別ごとに集計した。")

    table(doc,
        ["種別", "TRM要求数", "既存カバー済み", "未カバー", "未カバー率"],
        [
            ["BR（分岐網羅）", "55", "26", "29", "52.7%"],
            ["EC（同値クラス）", "27", "8", "19", "70.4%"],
            ["BV（境界値）", "11", "3", "8", "72.7%"],
            ["ER（エラーパス）", "3", "1", "2", "66.7%"],
            ["DP（依存切替）", "3", "0", "3", "100%"],
            ["合計", "99", "38", "61", "61.6%"],
        ],
        col_widths=[2.5, 2.0, 2.0, 2.0, 2.0], fs=8.5)

    para(doc,
        "全要求の61.6%（61件/99件）が既存テストでカバーされていなかった。"
        "特にDP（依存切替: 類似関数間の差異検証や往復変換検証）は"
        "既存テストに完全に欠落しており、BV（境界値）とEC（同値クラス）も"
        "70%以上が未カバーであった。"
        "一方、BR（分岐網羅）は比較的カバーされていた（未カバー率52.7%）。"
        "これは、既存の手動テストが主要な分岐パスの動作確認に偏り、"
        "境界値の体系的検証や関数間の整合性検証が不足していることを示す。")

    para(doc,
        "この分析は、TRMがテスト生成の「前」に品質ギャップを種別ごとに"
        "特定できることを実証するものであり、P2（網羅性の事前把握）への解決策となる。")

    # ------------------------------------------------------------------
    # 4. 結論
    # ------------------------------------------------------------------
    heading(doc, "4. 結論")

    para(doc,
        "本研究では、ソースコードの静的解析からテスト要求モデル（TRM）を"
        "YAML形式で生成し、TRMを入力としてLLMに単体テストを生成させる手法を"
        "提案し、OSSのC++プロジェクトの8関数（約385行）で実証評価した。"
        "得られた主要な結果は以下の4点である。")

    para(doc,
        "(1) TRM生成: 5種別（BR/EC/BV/ER/DP）で99件のテスト要求を構造化定義した。"
        "約385行のコードに対し平均3.9行/要求の密度でテスト要求を抽出した。"
        "これにより、テスト生成の「前」にテスト設計の全体像を把握・レビューできる"
        "状態を実現した（P1の解決）。")

    para(doc,
        "(2) テスト生成: TRMを入力として145件のテストコードを生成し、"
        "テスト要求カバー率100%を達成した。"
        "既存の手動テスト約45件の約3.2倍を体系的に生成した。")

    para(doc,
        "(3) 実行精度と障害診断: 初回PASS率91.7%。"
        "FAIL 12件は全てLLMの期待値推論誤りであり、"
        "TRM IDにより5つの原因分類に即時特定できた。"
        "修正後に全145件がPASSした（P3の解決）。")

    para(doc,
        "(4) 品質ギャップ発見: 既存テストの未カバー要求61件（61.6%）を"
        "種別ごとに特定した。特にDP種別は100%未カバーであり、"
        "BV・ECも70%以上が未カバーであった。"
        "TRMにより、テスト生成前に品質の弱点を構造的に把握できた（P2の解決）。")

    para(doc,
        "本実験は純粋関数8関数に限定されており、以下が今後の課題である。"
        "第一に、副作用を持つ関数への適用（スタブ・モック戦略の拡張が必要）。"
        "第二に、gcov/lcov等によるコードカバレッジ（C0/C1）との相関分析。"
        "第三に、他のプログラミング言語・プロジェクトでの再現実験。"
        "第四に、TRMの5種別体系のさらなる拡張"
        "（状態遷移テスト・組み合わせテスト等への対応）。")

    # 参考文献
    para(doc, "", sz=4, after=Pt(2))
    para(doc, "参考文献", bold=True, sz=9.5, after=Pt(3))

    refs = [
        # LLMテスト生成手法（比較対象）
        '[1] C. Lemieux, J. P. Inala, S. K. Lahiri, and S. Sen, '
        '"CodaMOSA: Escaping Coverage Plateaus in Test Generation with Pre-trained Large Language Models," '
        'in Proc. ICSE 2023, pp. 919-931, 2023.\n'
        'https://dl.acm.org/doi/10.1109/icse48619.2023.00085',

        '[2] Y. Chen et al., '
        '"ChatUniTest: A Framework for LLM-Based Test Generation," '
        'in Companion Proc. FSE 2024, ACM, 2024.\n'
        'https://dl.acm.org/doi/abs/10.1145/3663529.3663801',

        '[3] M. Schafer et al., '
        '"An Empirical Evaluation of Using Large Language Models for Automated Unit Test Generation," '
        'IEEE Transactions on Software Engineering, vol. 50, 2024.\n'
        'https://ieeexplore.ieee.org/document/10329992/',

        '[4] J. A. Ramos, E. M. Bader, J. Bader, and S. Iyer, '
        '"CoverUp: Coverage-Guided LLM-Based Test Generation," '
        'arXiv:2403.16218, 2024.\n'
        'https://arxiv.org/abs/2403.16218',

        # LLMテスト生成の品質課題
        '[5] N. S. Mathews and M. Nagappan, '
        '"Design choices made by LLM-based test generators prevent them from finding bugs," '
        'arXiv:2412.14137, 2024.\n'
        'https://arxiv.org/abs/2412.14137',

        '[6] B. Chu, Y. Feng, K. Liu, Z. Guo, Y. Zhang, H. Shi, Z. Nan, and B. Xu, '
        '"Large Language Models for Unit Test Generation: Achievements, Challenges, and the Road Ahead," '
        'arXiv:2511.21382, 2025.\n'
        'https://arxiv.org/abs/2511.21382',

        '[7] "Test smells in LLM-Generated Unit Tests," '
        'arXiv:2410.10628, 2024.\n'
        'https://arxiv.org/abs/2410.10628',

        '[8] "Do LLMs generate test oracles that capture the actual or the expected program behaviour?" '
        'arXiv:2410.21136, 2024.\n'
        'https://arxiv.org/abs/2410.21136',

        # 静的解析+LLM
        '[9] "Enhancing LLM-based Test Generation for Hard-to-Cover Branches via Program Analysis (TELPA)," '
        'ACM TOSEM, 2024.\n'
        'https://dl.acm.org/doi/10.1145/3748505',

        '[10] "LLMLOOP: Improving LLM-Generated Code and Tests through Automated Iterative Feedback Loops," '
        'in Proc. ICSME 2025, IEEE, 2025.\n'
        'https://arxiv.org/pdf/2603.23613',

        # 要求-テストトレーサビリティ・テスト設計
        '[11] "Multi-Step Generation of Test Specifications using Large Language Models," '
        'in Proc. ACL 2025 Industry Track, 2025.\n'
        'https://aclanthology.org/2025.acl-industry.11.pdf',

        '[12] G. Adzic, Specification by Example: How Successful Teams Deliver the Right Software, '
        'Manning, 2011.',

        # LLM生成コードの品質
        '[13] S. Dou et al., '
        '"What\'s Wrong with Your Code Generated by Large Language Models? An Extensive Study," '
        'arXiv:2407.06153, 2024.\n'
        'https://arxiv.org/abs/2407.06153',

        '[14] "Quality Assurance of LLM-generated Code: Addressing Non-Functional Quality Characteristics," '
        'arXiv:2511.10271, 2025.\n'
        'https://arxiv.org/html/2511.10271v1',

        # 日本のソフトウェア品質コミュニティ
        '[15] 「生成AIを活用したテストパターンマトリックスを用いたテスト生成」, '
        'SQiP 2024 一般発表 A4-1, 2024.\n'
        'https://www.juse.jp/sqip/symposium/2024/timetable/files/A4-1_ronbun.pdf',

        '[16] 産業技術総合研究所 (AIST), '
        '「生成AI品質マネジメントガイドライン第1版」, Rev. 1.0.0, 2025.\n'
        'https://www.digiarc.aist.go.jp/publication/aiqm/genaiqm-guidelines-v1.html',

        # マルチエージェント・ロードマップ
        '[17] "Large Language Models for Software Testing: A Research Roadmap," '
        'arXiv:2509.25043, 2025.\n'
        'https://arxiv.org/html/2509.25043v1',
    ]
    for ref in refs:
        para(doc, ref, sz=8, after=Pt(3))

    doc.save(OUT)
    print(f"Saved: {OUT}")
    print(f"File size: {os.path.getsize(OUT)} bytes")

if __name__ == "__main__":
    main()
