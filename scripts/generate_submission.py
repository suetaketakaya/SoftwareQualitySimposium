#!/usr/bin/env python3
"""SQiP 2026 提出用 docx 生成スクリプト"""

import os
import cairosvg
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE = "/Users/suetaketakaya/SoftwareQualitySymposium"
FIGURES = os.path.join(BASE, "drafts", "figures")
OUTPUT = os.path.join(BASE, "report", "submission_2026.docx")

# --- SVG -> PNG conversion ---
def convert_svgs():
    pngs = {}
    for name in ["fig1-overview", "fig2-requirements-breakdown",
                  "fig3-test-comparison", "fig4-execution-results"]:
        svg_path = os.path.join(FIGURES, f"{name}.svg")
        png_path = os.path.join(FIGURES, f"{name}.png")
        cairosvg.svg2png(url=svg_path, write_to=png_path, scale=2.0)
        pngs[name] = png_path
    return pngs

# --- Helper functions ---
def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_para(doc, text, bold=False, size=10.5, align=None, space_after=Pt(4)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = "游明朝"
    if align:
        p.alignment = align
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = Pt(0)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    # Rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

def add_figure(doc, png_path, caption, width=Cm(15)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(png_path, width=width)
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.font.size = Pt(9)
        run.italic = True
    cap.paragraph_format.space_after = Pt(6)
    return cap


# === MAIN ===
def main():
    pngs = convert_svgs()
    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    style = doc.styles["Normal"]
    style.font.name = "游明朝"
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.space_before = Pt(0)

    # ======================================================================
    # 発表申込書
    # ======================================================================
    add_para(doc, "ソフトウェア品質シンポジウム2026　「経験論文」「経験発表」", bold=True, size=12,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "発表申込／アブストラクト　記入用紙", bold=True, size=11,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))

    add_para(doc, "〔発表申込書〕", bold=True, size=11, space_after=Pt(8))

    add_para(doc, "タイトル", bold=True, size=10)
    add_para(doc, "テスト要求モデルを中間成果物とするLLM支援単体テスト生成手法"
             "\n── ホワイトボックステストの説明可能性と再現性の向上 ──", size=10.5)

    add_para(doc, "申込区分", bold=True, size=10)
    add_para(doc, "経験発表", size=10.5)

    add_para(doc, "カテゴリ", bold=True, size=10)
    add_para(doc, "・品質管理・テスト技術の観点", size=10.5)

    add_para(doc, "キーワード", bold=True, size=10)
    add_para(doc, "ホワイトボックステスト / テスト要求モデル / 大規模言語モデル（LLM） / "
             "静的解析 / 単体テスト自動化 / トレーサビリティ", size=10.5,
             space_after=Pt(12))

    # ======================================================================
    # Page break -> アブストラクト記入用紙
    # ======================================================================
    doc.add_page_break()
    add_para(doc, "〔アブストラクト記入用紙〕", bold=True, size=11,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(8))

    # ------------------------------------------------------------------
    # 1. ねらい
    # ------------------------------------------------------------------
    add_heading(doc, "1. ねらい", level=2)

    add_para(doc,
        "大規模言語モデル（LLM）によるテストコード自動生成は有望な技術であるが、"
        "LLMにソースコードを入力としてテストコードの生成を指示する直接生成アプローチには、"
        "実務上3つの課題がある。"
        "第一に説明可能性の欠如（生成テストの設計意図が不透明）、"
        "第二に再現性の不安定さ（同一プロンプトでも出力が変動）、"
        "第三に網羅性の不透明さ（カバー範囲が事後的にしか判断できない）である。"
        "これらは、LLMがテスト設計判断とコード記述を一体的に行い、"
        "設計意図が暗黙的にコードに埋め込まれることに起因する。")

    add_para(doc,
        "本研究は、ホワイトボックステスト要求モデルを人間が確認可能な中間成果物として"
        "明示的に生成・固定し、その要求に基づいてテスト生成を行う手法を提案する。"
        "これにより、(1) テスト設計の意図が可視化され説明可能性が高まる、"
        "(2) テスト要求モデルが固定されLLMのバージョンに依存しない再現性が得られる、"
        "(3) テスト要求のカバー率により網羅性を定量測定できる、という3つの効果を狙う。")

    add_para(doc,
        "先行研究としてCodaMosa [1]、ChatUniTest [2]、TestPilot [3] 等が"
        "LLMによるテスト生成を提案しているが、"
        "いずれもテストコードを直接生成するアプローチである。"
        "本研究の新規性は、テスト要求モデルという構造化された中間成果物を導入し、"
        "テスト設計判断（何をテストすべきか）とコード記述（どうテストするか）を"
        "明確に分離した点にある。")

    # ------------------------------------------------------------------
    # 2. 実施概要
    # ------------------------------------------------------------------
    add_heading(doc, "2. 実施概要", level=2)

    add_para(doc,
        "提案手法は3つのフェーズで構成される（図1）。", space_after=Pt(2))

    add_para(doc,
        "Phase 1（リポジトリ解析）: "
        "LLMが対象リポジトリのソースコード構造を解析し、"
        "純粋関数・明確な分岐構造・既存テスト存在・スタブ不要の4基準でテスト対象を選定する。"
        "人間がレビュー・承認する。")
    add_para(doc,
        "Phase 2（テスト要求モデル生成）: "
        "各関数のソースコードから分岐条件を抽出し、"
        "分岐網羅（BR）・同値クラス（EC）・境界値（BV）・エラーパス（ER）・"
        "依存切替（DP: 類似関数の差異検証や往復変換検証）の5種別で"
        "テスト要求をYAML形式で定義する。各要求にはID・説明・優先度・根拠を付与する。"
        "人間がレビュー・修正する。")
    add_para(doc,
        "Phase 3（テストコード生成）: "
        "承認されたテスト要求モデルを入力としてLLMがGoogle Test準拠のテストコードを生成する。"
        "各テストケースにテスト要求IDをコメントで明示し、トレーサビリティを確保する。")

    # Figure 1
    add_figure(doc, pngs["fig1-overview"], "図1: 提案手法の全体構成", width=Cm(15))

    add_para(doc,
        "実験設計: OSSのC++プロジェクト（Windows向け日本語テキストエディタ、"
        "Google Testによる既存テスト整備済み）に適用した。"
        "3領域8関数（合計約385行）を選定し、"
        "GUI依存・ファイルI/O依存等の11関数を除外理由を明示して除外した。"
        "比較の枠組みとして、(A) 手動テスト（既存テスト）、"
        "(B) LLMのみの直接生成（先行研究に基づく定性的評価）、"
        "(C) 提案方式の3方式を設定した。")

    # ------------------------------------------------------------------
    # 3. 実施結果
    # ------------------------------------------------------------------
    add_heading(doc, "3. 実施結果", level=2)

    add_para(doc,
        "Phase 2において8関数に対し計99件のテスト要求を定義した（図2）。"
        "種別内訳はBR 55件（55.6%）、EC 27件（27.3%）、BV 11件（11.1%）、"
        "ER 3件（3.0%）、DP 3件（3.0%）であった。"
        "Phase 3ではテスト要求モデルを入力としてテストコード計145件を生成し、"
        "テスト要求カバー率100%を達成した。"
        "なお、テスト要求カバー率100%はテスト要求の定義内容が漏れなくテストケースに"
        "反映されたことを意味し、テスト要求モデル自体の網羅性はC0/C1カバレッジとの"
        "相関分析で別途評価する予定である。")

    # Figure 2
    add_figure(doc, pngs["fig2-requirements-breakdown"],
               "図2: テスト要求の種別内訳（N=99）", width=Cm(10))

    add_para(doc,
        "生成テスト145件をmacOS環境（Apple clang 16.0 + Google Test 1.17 + "
        "Windows互換ヘッダ）でコンパイル・実行した。"
        "全3テスト実行ファイルがコンパイルに成功し、"
        "初回テストPASS率は91.7%（133/145件）であった。"
        "FAIL 12件の原因は全てLLMによるテスト期待値の推論誤りであり、"
        "対象関数の実装バグは0件であった（図4）。"
        "内訳は、ParseVersionのコンポーネント分割ロジック誤解5件、"
        "数字グルーピング誤解1件、IsMailAddressの文字数カウントミス2件・"
        "API呼び出しバグ1件、WhatKindOfTwoCharsのマッピング戻り値誤解3件であった。"
        "テスト要求モデルに照らして修正箇所を迅速に特定し、修正後に全145件がPASSした。"
        "この結果は、テスト要求モデルがテスト期待値の正誤判断基準としても機能することを示す。")

    # Figure 4
    add_figure(doc, pngs["fig4-execution-results"],
               "図4: コンパイル・実行結果", width=Cm(12))

    add_para(doc,
        "既存の手動テスト約45件に対し、提案手法では約3.2倍の145件を生成した（図3）。"
        "特に、修飾子の完全一致と部分一致の分岐（BR-20〜23）、"
        "有効文字判定の境界値（BV-07）、全ECharKind値の同種ペア網羅（EC-17）、"
        "通常版と4KW版の差異検証（DP-02）、往復変換検証（DP-03）など、"
        "既存テストで未カバーの要求を新規にカバーした。"
        "ただし、既存テストは回帰テストやバグ修正確認等の異なる目的で作成されたものであり、"
        "単純なテスト数の比較には留意が必要である。")

    # Figure 3
    add_figure(doc, pngs["fig3-test-comparison"],
               "図3: 領域別テストケース数の比較", width=Cm(12))

    # 三方式比較表
    add_para(doc, "表1: 三方式の定性的比較", bold=True, size=9.5,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))
    add_table(doc,
        ["評価軸", "(A) 手動テスト", "(B) LLMのみ (*)", "(C) 提案方式"],
        [
            ["テスト要求の明示性", "暗黙的", "なし", "明示的（YAML 99件）"],
            ["トレーサビリティ", "なし〜コメント", "なし", "ID対応で完全"],
            ["網羅性の事前測定", "カバレッジツール依存", "不可", "テスト要求カバー率で可能"],
            ["レビュー性", "テストコード直接", "テストコード直接", "要求+コードの二段階"],
            ["再現性", "担当者依存", "プロンプト依存", "要求モデルが固定"],
        ],
        col_widths=[3.5, 3.5, 3.5, 4.5])
    add_para(doc,
        "(*) 方式(B)は先行研究 [1-3] および著者の経験に基づく定性的評価であり、"
        "本実験での実測値ではない。", size=8.5, space_after=Pt(6))

    # ------------------------------------------------------------------
    # 4. 結論
    # ------------------------------------------------------------------
    add_heading(doc, "4. 結論", level=2)

    add_para(doc,
        "本研究では、静的解析で抽出したホワイトボックステスト要求を"
        "人間が確認可能な中間成果物として固定し、その要求に基づいてテスト生成を行う手法を提案した。"
        "OSSのC++プロジェクトの8関数（約385行）を対象に実証し、"
        "以下の結果を得た。"
        "(1) 5種別99件のテスト要求を体系的に定義し、テストコード生成の入力として機能することを確認した。"
        "(2) 生成テスト145件をコンパイル・実行し、初回91.7%がPASS、"
        "期待値修正後に100%がPASSした。12件の誤りはテスト要求モデルとの照合で迅速に修正できた。"
        "(3) 既存手動テスト約45件の約3.2倍を生成し、"
        "手動テストで見落とされやすいコーナーケースを新規にカバーした。"
        "(4) テスト要求IDによるトレーサビリティの確保とYAML形式の中間成果物により、"
        "ねらいで掲げた説明可能性・再現性・網羅性の3効果の実現を確認した。")

    add_para(doc,
        "テスト要求モデルの5種別はプログラミング言語に依存しない汎用的な分類体系であり、"
        "他言語・他プロジェクトへの適用の方向性を示唆する。"
        "ただし、本実験はC++の単一プロジェクト・純粋関数に限定されており、"
        "副作用を持つ関数への拡張（モック・スタブの自動生成）、"
        "C0/C1カバレッジの実測、方式(B)との定量比較、"
        "他言語での追試が今後の課題である。")

    # 参考文献
    add_para(doc, "", size=6, space_after=Pt(2))
    add_para(doc, "参考文献", bold=True, size=9)
    add_para(doc,
        "[1] Lemieux, C., et al. \"CodaMosa: Escaping coverage plateaus in test "
        "generation with pre-trained large language models.\" ICSE 2023.\n"
        "[2] Chen, Y., et al. \"ChatUniTest: A Framework for LLM-Based Test "
        "Generation.\" FSE 2024.\n"
        "[3] Schäfer, M., et al. \"An Empirical Evaluation of Using Large Language "
        "Models for Automated Unit Test Generation.\" IEEE TSE 2024.",
        size=8.5, space_after=Pt(0))

    # --- Save ---
    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")

if __name__ == "__main__":
    main()
